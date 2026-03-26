import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 标题
title = doc.add_heading('AHL商务模式评估报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('五种定价模式全面对比与战略建议', level=3)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('项目: AHL (AI Hotel Language)')
doc.add_paragraph('版本: V1.0 评估报告')
doc.add_paragraph('日期: 2026年3月23日')
doc.add_paragraph('')

# 一、五种定价模式定义
doc.add_heading('一、五种定价模式定义', level=1)

doc.add_heading('1.1 模式总览', level=2)
table = doc.add_table(rows=7, cols=5)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '模式'
hdr[1].text = '初装费'
hdr[2].text = '月度订阅费'
hdr[3].text = 'TOKEN定价'
hdr[4].text = '核心特点'
data = [
    ('模式1', '¥15K-50K', '¥300-5,000', '含在月费', '传统SaaS'),
    ('模式2', '¥0', '¥300-5,000', '阶梯定价', '混合型阶梯'),
    ('模式3', '¥0', '¥0', '阶梯定价', '纯TOKEN阶梯'),
    ('模式4', '¥0', '¥300-5,000', '统一价格', '混合型特斯拉'),
    ('模式5', '¥0', '¥0', '统一价格', '纯TOKEN统一')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

# 二、底层逻辑分析
doc.add_heading('二、底层逻辑分析', level=1)

doc.add_heading('2.1 收费原理', level=2)
doc.add_paragraph('为什么收费？')
doc.add_paragraph('- AHL提供的是AI能力，本质是大模型TOKEN消耗')
doc.add_paragraph('- TOKEN是AI运行的"燃料"，按量收费是行业趋势')
doc.add_paragraph('- 酒店使用AI越多，价值创造越大，收费越合理')

doc.add_heading('2.2 B端商家接受度分析', level=2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '费用类型'
hdr[1].text = '酒店接受度'
hdr[2].text = '原因'
data = [
    ('初装费', '❌ 厌恶', '倾向零风险尝试'),
    ('月度订阅费', '✅ 理解', '类比传统SaaS'),
    ('TOKEN计费', '⚠️ 需教育', '理解"用多少付多少"')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('2.3 特斯拉模式分析', level=2)
doc.add_paragraph('特斯拉定价模式核心逻辑：')
doc.add_paragraph('- 产品过硬 → 价值透明 → 统一价格')
doc.add_paragraph('- 不搞阶梯、不砍价 → 简单高效')
doc.add_paragraph('- 降低交易成本 → 专注产品')
doc.add_paragraph('')
doc.add_paragraph('AHL借鉴：')
doc.add_paragraph('- TOKEN费采用统一单价（特斯拉模式）')
doc.add_paragraph('- 简单透明，减少客户决策成本')

# 三、市场现状分析
doc.add_heading('三、市场现状分析', level=1)

doc.add_heading('3.1 AI SaaS市场现状', level=2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '模式'
hdr[1].text = '代表案例'
hdr[2].text = '特点'
data = [
    ('纯订阅费', 'Salesforce, HubSpot', '稳定但增长慢'),
    ('纯TOKEN', 'OpenAI, Anthropic', '增长快但波动大'),
    ('混合模式', 'Microsoft Copilot', '平衡但复杂')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('3.2 酒店行业SaaS现状', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '类型'
hdr[1].text = '代表产品'
hdr[2].text = '收费模式'
data = [
    ('PMS', '石基、Opera', '初装费+年费'),
    ('CRS', '直采、厦门', '年费'),
    ('OTA运营工具', '众荟、声誉', '月费/年费'),
    ('收益管理', 'IDeaS, Duetto', '按房间数订阅')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

# 四、五种模式全面对比
doc.add_heading('四、五种模式全面对比', level=1)

doc.add_heading('4.1 核心指标对比', level=2)
table = doc.add_table(rows=8, cols=6)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '指标'
hdr[1].text = '模式1'
hdr[2].text = '模式2'
hdr[3].text = '模式3'
hdr[4].text = '模式4'
hdr[5].text = '模式5'
data = [
    ('B端接受度', '⭐⭐⭐⭐', '⭐⭐⭐⭐', '⭐⭐', '⭐⭐⭐⭐⭐', '⭐⭐⭐'),
    ('规模化速度', '⭐⭐', '⭐⭐⭐', '⭐⭐⭐⭐⭐', '⭐⭐⭐⭐', '⭐⭐⭐⭐⭐'),
    ('弱运营程度', '⭐⭐', '⭐⭐⭐', '⭐⭐⭐⭐⭐', '⭐⭐⭐⭐', '⭐⭐⭐⭐⭐'),
    ('低成本结构', '⭐⭐', '⭐⭐⭐', '⭐⭐⭐⭐', '⭐⭐⭐⭐', '⭐⭐⭐⭐⭐'),
    ('高效率运转', '⭐⭐', '⭐⭐⭐', '⭐⭐⭐⭐⭐', '⭐⭐⭐⭐', '⭐⭐⭐⭐⭐'),
    ('短期收入保障', '⭐⭐⭐⭐⭐', '⭐⭐⭐⭐', '⭐⭐', '⭐⭐⭐⭐', '⭐⭐'),
    ('长期增长潜力', '⭐⭐', '⭐⭐⭐', '⭐⭐⭐⭐⭐', '⭐⭐⭐⭐', '⭐⭐⭐⭐⭐')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('4.2 收入预测对比（36个月）', level=2)
table = doc.add_table(rows=7, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '模式'
hdr[1].text = '收入特征'
hdr[2].text = '36个月累计'
hdr[3].text = '毛利率'
data = [
    ('模式1', '初装+月费', '¥2.85亿', '52%'),
    ('模式2', '免装+月费+阶梯', '¥4.5亿', '35%'),
    ('模式3', '免装+免月+阶梯', '¥8.5亿', '42%'),
    ('模式4', '免装+月费+统一', '¥5.8亿', '38%'),
    ('模式5', '免装+免月+统一', '¥11.2亿', '46%'),
    ('模式4*', '免装+月费+统一(推荐)', '¥5.8亿', '38%')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

# 五、战略评估
doc.add_heading('五、战略评估', level=1)

doc.add_heading('5.1 核心矛盾分析', level=2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '矛盾'
hdr[1].text = '解读'
hdr[2].text = '平衡点'
data = [
    ('B端接受度 vs 规模化', '初装/月费提高接受度但降低速度', '免初装是必须的'),
    ('运营复杂度 vs 效率', '阶梯定价复杂但符合商业逻辑', '统一定价更简单'),
    ('短期收入 vs 长期规模', '月费保短期但影响规模', '渐进式过渡')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('5.2 目标匹配度', level=2)
table = doc.add_table(rows=6, cols=6)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '目标'
hdr[1].text = '模式1'
hdr[2].text = '模式2'
hdr[3].text = '模式3'
hdr[4].text = '模式4'
hdr[5].text = '模式5'
data = [
    ('B端接受度', '✅', '✅', '❌', '✅✅', '✅'),
    ('快速规模化', '❌', '✅', '✅✅', '✅', '✅✅'),
    ('弱运营', '❌', '✅', '✅✅', '✅', '✅✅'),
    ('低成本', '❌', '✅', '✅', '✅', '✅✅'),
    ('高效率', '❌', '✅', '✅✅', '✅', '✅✅')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

# 六、模式4深度分析（推荐）
doc.add_heading('六、模式4深度分析（推荐）', level=1)

doc.add_heading('6.1 模式4核心逻辑', level=2)
doc.add_paragraph('免初装 + 月度订阅 + TOKEN统一定价（特斯拉模式）')
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '费用类型'
hdr[1].text = '定价依据'
hdr[2].text = '说明'
data = [
    ('初装费', '¥0', '降低入驻门槛'),
    ('月度订阅费', '按房间数', '¥300-5,000/月'),
    ('TOKEN费', '统一单价', '¥0.005/Token')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('6.2 月度订阅费（按房间数）', level=2)
table = doc.add_table(rows=8, cols=2)
table.style = 'Table Grid'
data = [
    ('房间数量', '月度订阅费'),
    ('50间以下', '¥300/月'),
    ('50-100间', '¥500/月'),
    ('100-150间', '¥800/月'),
    ('150-200间', '¥1,200/月'),
    ('200-300间', '¥1,800/月'),
    ('300-500间', '¥2,800/月'),
    ('500间以上', '¥5,000/月')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

doc.add_heading('6.3 TOKEN费（统一单价）', level=2)
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
data = [
    ('项目', '价格'),
    ('统一单价', '¥0.005/Token'),
    ('计费方式', '按实际消耗')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

doc.add_heading('6.4 模式4综合定价示例', level=2)
table = doc.add_table(rows=7, cols=5)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '酒店类型'
hdr[1].text = '房间数'
hdr[2].text = '月度订阅'
hdr[3].text = 'TOKEN月费'
hdr[4].text = '月均总费用'
data = [
    ('奢华酒店', '200间', '¥1,800', '¥1,000-1,500', '¥2,800-3,300'),
    ('高端酒店', '200间', '¥1,800', '¥800-1,200', '¥2,600-3,000'),
    ('中高端酒店', '150间', '¥800', '¥400-700', '¥1,200-1,500'),
    ('中端酒店', '100间', '¥500', '¥250-400', '¥750-900'),
    ('经济型酒店', '80间', '¥300', '¥150-250', '¥450-550'),
    ('民宿', '20间', '¥300', '¥100-200', '¥400-500')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

# 七、渐进式发展路径
doc.add_heading('七、渐进式发展路径', level=1)

doc.add_heading('7.1 阶段一：种子期（M1-M6）', level=2)
doc.add_paragraph('目标: 50-200家商户')
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'
data = [
    ('策略', '月费5折 + TOKEN5折'),
    ('目的', '快速获取种子用户')
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_heading('7.2 阶段二：验证期（M7-M12）', level=2)
doc.add_paragraph('目标: 500-1000家商户')
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'
data = [
    ('策略', '月费8折 + TOKEN9折'),
    ('目的', '验证商业模式')
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_heading('7.3 阶段三：增长期（M13-M24）', level=2)
doc.add_paragraph('目标: 2000-4000家商户')
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'
data = [
    ('策略', '全价执行'),
    ('目的', '规模化扩张')
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_heading('7.4 阶段四：过渡期（M25-M36）', level=2)
doc.add_paragraph('目标: 5000+家商户')
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'
data = [
    ('策略', '逐步降低月费'),
    ('目的', '向模式5过渡')
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_heading('7.5 阶段五：成熟期（M36+）', level=2)
doc.add_paragraph('目标: 行业标准')
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'
data = [
    ('策略', '纯TOKEN为主'),
    ('目的', '模式5最终形态')
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

# 八、最终建议
doc.add_heading('八、最终建议', level=1)

doc.add_heading('8.1 推荐方案', level=2)
doc.add_paragraph('')
recommendation = doc.add_paragraph()
recommendation.add_run('🎯 首选：模式4（免初装+月度订阅+TOKEN统一定价）').bold = True

doc.add_paragraph('')
doc.add_paragraph('核心理由：')
doc.add_paragraph('1. B端接受度最高 - 免初装+月费习惯+TOKEN统一')
doc.add_paragraph('2. 规模化与收入平衡 - 短期月费+长期TOKEN')
doc.add_paragraph('3. 弱运营/低成本 - 统一定价+自动计费')
doc.add_paragraph('4. 符合AI发展趋势 - TOKEN是未来')

doc.add_heading('8.2 模式对比总结', level=2)
table = doc.add_table(rows=7, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '模式'
hdr[1].text = '优势'
hdr[2].text = '劣势'
hdr[3].text = '适用场景'
data = [
    ('模式1', '收入稳定', '规模化慢', '传统酒店集团'),
    ('模式2', '平衡', '复杂', '中期过渡'),
    ('模式3', '易规模化', '收入风险', '资本充足时'),
    ('模式4*', '平衡+简单', '需过渡', '首选'),
    ('模式5', '最简', '收入不确定', '长期目标')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('8.3 核心洞察', level=2)
doc.add_paragraph('1. 免初装是必须的：降低入驻门槛是快速规模化的前提')
doc.add_paragraph('2. 特斯拉模式是对的：统一价格简单透明，符合产品自信逻辑')
doc.add_paragraph('3. 月费是过渡桥：短期保障收入，中期逐步降低')
doc.add_paragraph('4. TOKEN是未来：长期看，纯TOKEN模式是最终形态')

# 九、风险提示
doc.add_heading('九、风险提示', level=1)

doc.add_heading('9.1 风险因素', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '风险'
hdr[1].text = '影响'
hdr[2].text = '对策'
data = [
    ('大模型成本波动', 'TOKEN毛利不确定', '动态调整定价'),
    ('酒店预算削减', '月费收入下降', '保持灵活性'),
    ('竞争加剧', '价格战风险', '差异化价值'),
    ('市场教育不足', 'TOKEN接受度低', '渐进式推进')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('9.2 应对策略', level=2)
doc.add_paragraph('- 保持定价灵活性')
doc.add_paragraph('- 持续提升产品价值')
doc.add_paragraph('- 建立客户成功团队')
doc.add_paragraph('- 积累数据和案例')

# 保存
desktop = os.path.join(os.environ['USERPROFILE'], 'Desktop')
output_path = os.path.join(desktop, 'AHL-商务模式评估报告V1.0.docx')
doc.save(output_path)
print(f'已保存到: {output_path}')
