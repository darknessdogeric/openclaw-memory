#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convert Markdown report to Word document
"""

import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='SimSun', size=10.5, bold=False):
    """Set Chinese font for a run"""
    font = run.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_zh(doc, text, level=1):
    """Add Chinese heading"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    font_names = ['SimHei', 'Microsoft YaHei', 'SimSun']
    font_size = {1: 18, 2: 16, 3: 14, 4: 12}.get(level, 12)
    set_chinese_font(run, font_name=font_names[0], size=font_size, bold=True)
    return heading

def add_paragraph_zh(doc, text, bold=False, size=10.5, alignment=None):
    """Add Chinese paragraph"""
    para = doc.add_paragraph()
    if alignment:
        para.alignment = alignment
    run = para.add_run(text)
    set_chinese_font(run, font_name='SimSun', size=size, bold=bold)
    return para

def parse_markdown_table(lines, start_idx):
    """Parse markdown table and return data and end index"""
    data = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        line = lines[i].strip()
        # Skip separator line
        if '---' in line or line.replace('|', '').replace('-', '').replace(' ', '') == '':
            i += 1
            continue
        # Parse cells
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            data.append(cells)
        i += 1
    return data, i

def markdown_to_docx(md_file, docx_file):
    """Convert markdown file to Word document"""
    
    # Create document
    doc = Document()
    
    # Set default font for document
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(10.5)
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_table = False
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Parse headings
        if line.startswith('# '):
            add_heading_zh(doc, line[2:], level=1)
            i += 1
            continue
        elif line.startswith('## '):
            add_heading_zh(doc, line[3:], level=2)
            i += 1
            continue
        elif line.startswith('### '):
            add_heading_zh(doc, line[4:], level=3)
            i += 1
            continue
        elif line.startswith('#### '):
            add_heading_zh(doc, line[5:], level=4)
            i += 1
            continue
        
        # Parse tables
        if line.startswith('|') and not in_table:
            table_data, end_idx = parse_markdown_table(lines, i)
            if len(table_data) >= 1:
                # Create table
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                
                for row_idx, row_data in enumerate(table_data):
                    row = table.rows[row_idx]
                    for col_idx, cell_text in enumerate(row_data):
                        cell = row.cells[col_idx]
                        cell.text = cell_text
                        # Set font for cell
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                set_chinese_font(run, font_name='SimSun', size=9)
                
                # Format header row
                if len(table_data) > 0:
                    for cell in table.rows[0].cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                set_chinese_font(run, font_name='SimHei', size=9, bold=True)
            
            i = end_idx
            continue
        
        # Parse blockquotes
        if line.startswith('>'):
            text = line[1:].strip()
            para = add_paragraph_zh(doc, text, bold=False, size=10)
            para.paragraph_format.left_indent = Inches(0.3)
            i += 1
            continue
        
        # Parse list items
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            para = add_paragraph_zh(doc, '• ' + text, bold=False, size=10.5)
            para.paragraph_format.left_indent = Inches(0.3)
            i += 1
            continue
        
        # Parse numbered lists
        if line[0:2].replace('.', '').isdigit():
            text = line[line.find(' ')+1:]
            para = add_paragraph_zh(doc, line[0:line.find(' ')+1] + text, bold=False, size=10.5)
            para.paragraph_format.left_indent = Inches(0.3)
            i += 1
            continue
        
        # Parse code blocks
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_text = '\n'.join(code_lines)
            para = add_paragraph_zh(doc, code_text, bold=False, size=9)
            para.paragraph_format.left_indent = Inches(0.3)
            i += 1
            continue
        
        # Parse inline code
        if '`' in line:
            parts = line.split('`')
            para = doc.add_paragraph()
            for idx, part in enumerate(parts):
                run = para.add_run(part)
                if idx % 2 == 1:  # Code part
                    set_chinese_font(run, font_name='Courier New', size=9, bold=True)
                else:
                    set_chinese_font(run, font_name='SimSun', size=10.5)
            i += 1
            continue
        
        # Parse bold text
        if '**' in line:
            parts = line.split('**')
            para = doc.add_paragraph()
            for idx, part in enumerate(parts):
                if idx % 2 == 1:  # Bold part
                    run = para.add_run(part)
                    set_chinese_font(run, font_name='SimSun', size=10.5, bold=True)
                else:
                    run = para.add_run(part)
                    set_chinese_font(run, font_name='SimSun', size=10.5)
            i += 1
            continue
        
        # Regular paragraph
        add_paragraph_zh(doc, line, bold=False, size=10.5)
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"Document saved to: {docx_file}")

if __name__ == '__main__':
    if len(sys.argv) != 3:
        print("Usage: python md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    
    markdown_to_docx(sys.argv[1], sys.argv[2])
