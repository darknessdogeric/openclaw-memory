from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 标题
title = doc.add_heading('AHL商业模式对比分析方案', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 副标题
subtitle = doc.add_paragraph()
subtitle.add_run('三种定价策略全方位对比与预测').bold = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('项目: AHL (AI Hotel Language)')
doc.add_paragraph('版本: V1.0')
doc.add_paragraph('日期: 2026年3月23日')
doc.add_paragraph('')

# 一、三种商业模式定义
doc.add_heading('一、三种商业模式定义', level=1)

doc.add_heading('方案A：初装费 + 月度维护费（传统SaaS模式）', level=2)
table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
cells = table.rows[0].cells
cells[0].text = '项目'
cells[1].text = '定价'
data = [
    ('初装费', '¥15,000-30,000/酒店'),
    ('月度维护费', '¥800-2,000/月'),
    ('交易抽成', '0%')
]
for i, (k, v) in enumerate(data):
    cells = table.rows[i+1].cells
    cells[0].text = k
    cells[1].text = v

doc.add_heading('方案B：免初装 + 月度维护 + TOKEN计费（混合模式）', level=2)
table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
cells = table.rows[0].cells
cells[0].text = '项目'
cells[1].text = '定价'
data = [
    ('初装费', '¥0（免费部署）'),
    ('月度维护费', '¥300-800/月'),
    ('TOKEN计费', '¥0.003-0.008/Token')
]
for i, (k, v) in enumerate(data):
    cells = table.rows[i+1].cells
    cells[0].text = k
    cells[1].text = v

doc.add_heading('方案C：免初装 + 免维护 + TOKEN计费（AI原生模式）', level=2)
table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'
cells = table.rows[0].cells
cells[0].text = '项目'
cells[1].text = '定价'
data = [
    ('初装费', '¥0（完全免费）'),
    ('月度维护费', '¥0'),
    ('TOKEN计费', '¥0.004-0.010/Token')
]
for i, (k, v) in enumerate(data):
    cells = table.rows[i+1].cells
    cells[0].text = k
    cells[1].text = v

# 二、核心假设参数
doc.add_heading('二、核心假设参数', level=1)

doc.add_heading('2.1 市场假设', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '参数'
hdr[1].text = '数值'
hdr[2].text = '说明'
data = [
    ('目标商户总数', '5,000家', '3年期目标'),
    ('年均GMV/商户', '¥200万', '中小型酒店平均水平'),
    ('平台交易费率', '2.5%', 'AHL标准费率'),
    ('单房年均收入', '¥8万', '100间房酒店')
]
for i, (k, v, d) in enumerate(data):
    cells = table.rows[i+1].cells
    cells[0].text = k
    cells[1].text = v
    cells[2].text = d

doc.add_heading('2.2 TOKEN消耗假设', level=2)
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '酒店规模'
hdr[1].text = '日均TOKEN'
hdr[2].text = '月均TOKEN'
hdr[3].text = '月均成本'
data = [
    ('50间以下', '50,000', '1.5M', '¥3,000-7,500'),
    ('50-100间', '120,000', '3.6M', '¥7,200-18,000'),
    ('100-200间', '250,000', '7.5M', '¥15,000-37,500'),
    ('200间以上', '500,000', '15M', '¥30,000-75,000')
]
for i, (a, b, c, d) in enumerate(data):
    cells = table.rows[i+1].cells
    cells[0].text = a
    cells[1].text = b
    cells[2].text = c
    cells[3].text = d

# 三、收入预测
doc.add_heading('三、收入预测模型', level=1)

doc.add_heading('3.1 方案A：初装费 + 月度维护费', level=2)
table = doc.add_table(rows=7, cols=6)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '月份'
hdr[1].text = '新签'
hdr[2].text = '累计'
hdr[3].text = '初装收入'
hdr[4].text = '月费收入'
hdr[5].text = '总收入'
data = [
    ('M1', '15', '15', '¥300K', '¥12K', '¥312K'),
    ('M6', '50', '195', '¥900K', '¥234K', '¥1,134K'),
    ('M12', '120', '915', '¥2,400K', '¥1,098K', '¥3,498K'),
    ('M18', '200', '2,415', '¥4,000K', '¥3,218K', '¥7,218K'),
    ('M24', '250', '3,915', '¥5,000K', '¥5,510K', '¥10,510K'),
    ('M36', '90', '5,000', '¥1,800K', '¥7,500K', '¥9,300K')
]
for i, row in enumerate(data):
    cells = table.rows[i+1].cells
    for j, val in enumerate(row):
        cells[j].text = val

p = doc.add_paragraph()
p.add_run('36个月累计收入: ¥2.85亿').bold = True

doc.add_heading('3.2 方案B：免初装 + 月度维护 + TOKEN', level=2)
table = doc.add_table(rows=7, cols=5)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '月份'
hdr[1].text = '累计'
hdr[2].text = '月费收入'
hdr[3].text = 'TOKEN收入'
hdr[4].text = '总收入'
data = [
    ('M1', '15', '¥7.5K', '¥22.5K', '¥30K'),
    ('M6', '195', '¥97.5K', '¥585K', '¥682K'),
    ('M12', '915', '¥457.5K', '¥4,575K', '¥5,033K'),
    ('M18', '2,415', '¥1,207K', '¥18,112K', '¥19,319K'),
    ('M24', '3,915', '¥1,957K', '¥39,150K', '¥41,107K'),
    ('M36', '5,000', '¥2,500K', '¥62,500K', '¥65,000K')
]
for i, row in enumerate(data):
    cells = table.rows[i+1].cells
    for j, val in enumerate(row):
        cells[j].text = val

p = doc.add_paragraph()
p.add_run('36个月累计收入: ¥5.8亿').bold = True

doc.add_heading('3.3 方案C：免初装 + 免维护 + TOKEN', level=2)
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '月份'
hdr[1].text = '累计'
hdr[2].text = 'TOKEN收入'
data = [
    ('M1', '15', '¥45K'),
    ('M6', '195', '¥1,170K'),
    ('M12', '915', '¥9,150K'),
    ('M18', '2,415', '¥36,225K'),
    ('M24', '3,915', '¥78,300K'),
    ('M36', '5,000', '¥125,000K')
]
for i, row in enumerate(data):
    cells = table.rows[i+1].cells
    for j, val in enumerate(row):
        cells[j].text = val

p = doc.add_paragraph()
p.add_run('36个月累计收入: ¥11.2亿').bold = True

# 四、综合对比
doc.add_heading('四、三方案综合对比', level=1)

doc.add_heading('4.1 核心指标对比', level=2)
table = doc.add_table(rows=11, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '指标'
hdr[1].text = '方案A'
hdr[2].text = '方案B'
hdr[3].text = '方案C'
data = [
    ('模式', '初装+月费', '免装+月费+TOKEN', '免装+免维护+TOKEN'),
    ('36个月累计收入', '¥2.85亿', '¥5.8亿', '¥11.2亿'),
    ('月均收入', '¥79万', '¥161万', '¥311万'),
    ('毛利率', '52%', '16%', '46%'),
    ('累计毛利', '¥1.47亿', '¥0.91亿', '¥5.1亿'),
    ('收入增长曲线', '线性', '指数', '指数'),
    ('商户入驻门槛', '高', '中', '零'),
    ('市场扩张速度', '慢', '中', '快'),
    ('可参照模式', '传统SaaS', 'Mixpanel式', 'ChatGPT式')
]
for i, row in enumerate(data):
    cells = table.rows[i+1].cells
    for j, val in enumerate(row):
        cells[j].text = val

# 五、建议
doc.add_heading('五、综合建议', level=1)

doc.add_paragraph('')
recommendation = doc.add_paragraph()
recommendation.add_run('首选方案：方案B（混合模式）').bold = True

reasons = doc.add_paragraph('理由：')
reasons.add_run('1. 现金流有保障 - 月度维护费是"保底"\n')
reasons.add_run('2. 市场扩张快 - 零初装降低入驻门槛\n')
reasons.add_run('3. 增长可期待 - TOKEN收入弹性大\n')
reasons.add_run('4. 符合AI趋势 - 不与产业发展对抗')

doc.add_paragraph('')
path = doc.add_paragraph('渐进式路径：')
path.add_run('M1-M12: 方案B (月费¥500 + TOKEN¥0.005)\n')
path.add_run('M13-M24: 方案B过渡 (月费¥200 + TOKEN¥0.006)\n')
path.add_run('M25+: 可选方案C (纯TOKEN ¥0.007)')

doc.add_paragraph('')
conclusion = doc.add_paragraph()
conclusion.add_run('核心洞察: TOKEN是"过路费"，可以收一辈子；初装费是"门票"，只能收一次。过路费的规模远大于门票。').bold = True

# 保存
import os
desktop = os.path.join(os.environ['USERPROFILE'], 'Desktop')
output_path = os.path.join(desktop, 'AHL-商业模式对比分析-V1.0.docx')
doc.save(output_path)
print(f'已保存到: {output_path}')
