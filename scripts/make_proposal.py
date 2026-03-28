# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

def set_run_font(run, font_name='微软雅黑', size=11, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)

def shade_cell(cell, fill):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def header_row(table, texts, fill='2E75B6'):
    cells = table.rows[0].cells
    for i, t in enumerate(texts):
        cells[i].text = t
        for p in cells[i].paragraphs:
            for run in p.runs:
                set_run_font(run, '微软雅黑', 10, bold=True)
                run.font.color.rgb = RGBColor(255, 255, 255)
        shade_cell(cells[i], fill)

# ===== 标题 =====
t = doc.add_heading('AHL 酒店智能升级项目', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in t.runs:
    set_run_font(run, '微软雅黑', 22, bold=True)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('合 作 意 向 书')
set_run_font(run, '微软雅黑', 16, bold=True)
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('版本 V1.0 | 2026年3月 | 免费试点合作邀请')
set_run_font(run, '微软雅黑', 9, color=RGBColor(0x80, 0x80, 0x80))

doc.add_paragraph()

# ===== 一封邀请函 =====
h = doc.add_heading('一封邀请函', 2)
for run in h.runs:
    set_run_font(run, '微软雅黑', 14, bold=True)

p = doc.add_paragraph()
run = p.add_run('尊敬的 [酒店名称] 管理团队 / 业主方：')
set_run_font(run, '微软雅黑', 11, bold=True)

p = doc.add_paragraph()
run = p.add_run(
    '我们正在建设一个覆盖全国酒店的AI智能服务网络，邀请贵酒店作为'
    '首批试点合作伙伴，免费体验为期【X个月】的AI数字员工服务。'
)
set_run_font(run, '微软雅黑', 11)

p = doc.add_paragraph()
run = p.add_run(
    '不需要任何投入，我们带着技术、工具和团队来驻点，目标是：让贵酒店的'
    '线上获客能力、私域运营效率、宾客满意度，在试点期内看到可量化的提升。'
)
set_run_font(run, '微软雅黑', 11)

p = doc.add_paragraph()
run = p.add_run(
    '如果效果达到预期，我们期待与贵酒店建立长期合作关系。'
    '如果效果未达预期，试点期结束后贵酒店可以选择退出，不产生任何费用或遗留问题。'
)
set_run_font(run, '微软雅黑', 11, bold=True, color=RGBColor(0xC0, 0x00, 0x00))

doc.add_paragraph()

# ===== 行业机会 =====
h = doc.add_heading('写在前面：我们观察到的一个行业机会', 2)
for run in h.runs:
    set_run_font(run, '微软雅黑', 14, bold=True)

p = doc.add_paragraph()
run = p.add_run('过去三年，酒店的线上获客成本在持续上涨。')
set_run_font(run, '微软雅黑', 11)

items = [
    ('OTA佣金压力：', '携程、美团的佣金比率通常在10%-15%之间，100元客房收入有10-15元交给平台'),
    ('数据不在自己手里：', '客人的数据在平台上，不在酒店手里。酒店努力做口碑，但复购客人还是先去携程搜'),
    ('客人习惯变了：', '越来越少人愿意打电线订房，越来越多人希望在微信里直接完成预订'),
]
for t2, d in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(t2)
    set_run_font(run, '微软雅黑', 11, bold=True)
    run = p.add_run(d)
    set_run_font(run, '微软雅黑', 11)

p = doc.add_paragraph()
run = p.add_run('供需两端都在变，但大多数酒店的运营方式，还停留在五年前。这就是我们看到的机会。')
set_run_font(run, '微软雅黑', 11, bold=True)

doc.add_paragraph()

# ===== 这个项目是什么 =====
h = doc.add_heading('这个项目是什么', 2)
for run in h.runs:
    set_run_font(run, '微软雅黑', 14, bold=True)

p = doc.add_paragraph()
run = p.add_run('AHL（AI Hotel Language）是一个酒店AI智能服务平台，核心理念是：')
set_run_font(run, '微软雅黑', 11)

q = doc.add_paragraph()
q.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = q.add_run('"让每一家酒店，都拥有自己的AI运营团队"')
set_run_font(run, '微软雅黑', 13, bold=True, color=RGBColor(0x2E, 0x75, 0xB6))

p = doc.add_paragraph()
run = p.add_run('我们把AI能力封装成5大模块，可以单独启用，也可以组合使用：')
set_run_font(run, '微软雅黑', 11, bold=True)

doc.add_paragraph()

# 5大能力模块标题
p = doc.add_paragraph()
run = p.add_run('🎯 5大AI能力模块')
set_run_font(run, '微软雅黑', 12, bold=True, color=RGBColor(0xC0, 0x00, 0x00))

# 5大模块表格
tbl = doc.add_table(rows=6, cols=3)
tbl.style = 'Table Grid'
header_row(tbl, ['模块', '解决什么问题', '给酒店带来什么'], '2E75B6')

modules = [
    ('① 收益管理与自动调价', '定价靠经验、靠感觉，跟不上市场变化', 'AI实时分析竞品/需求/天气，自动给出调价建议，收益人员一键确认执行'),
    ('② OTA智能维护', '携程评分掉了一片还不知道、差评48h才发现、回复差评靠抄模板', '差评30分钟内预警到手机，AI生成个性化回复建议，紧急差评立即电话通知'),
    ('③ 私域运营自动化', '客人离店后就像"失联"一样，再也没来过', '离店自动入微信社群，定时推送天气/活动/优惠，复购触达不用人工干预'),
    ('④ C端：客人自然语言预订与服务', '客人还要打电线才能订房，问个问题要等前台回复', '客人在微信里说"周六带孩子住江景600左右"，AI 30秒内推荐最合适的房型，协助完成预订'),
    ('⑤ B端：酒店自然语言自动运营', '酒店要查数据/做报表/查竞品，得找IT或手工统计', '管理者问"本周入住率多少？""周边竞品价格有什么变化？"AI直接给出答案和解读'),
]
for i, (name, prob, sol) in enumerate(modules):
    row = tbl.rows[i + 1]
    row.cells[0].text = name
    row.cells[1].text = prob
    row.cells[2].text = sol
    shade_cell(row.cells[0], 'D6E4F0')
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, '微软雅黑', 9)

doc.add_paragraph()

# ===== 传统 vs AHL ======
h = doc.add_heading('传统模式 vs AHL模式', 2)
for run in h.runs:
    set_run_font(run, '微软雅黑', 14, bold=True)

tbl2 = doc.add_table(rows=6, cols=2)
tbl2.style = 'Table Grid'
header_row(tbl2, ['传统模式', 'AHL模式'])

comp_rows = [
    ('携程一条独大，佣金10-15%', '直客通私域预订，降低OTA依赖'),
    ('客人打电话问房型，前台一个个回', '客人微信说需求，AI 30秒内接待'),
    ('差评出现后48小时才发现', '差评出现后30分钟内预警到店长手机'),
    ('收益管理靠经验，拍脑袋定价', 'AI实时分析竞品数据，给出动态调价建议'),
    ('离店客人石沉大海，再也没有复购', '离店自动入社群，持续触达激活复购'),
]
for i, (old, new) in enumerate(comp_rows):
    row = tbl2.rows[i + 1]
    row.cells[0].text = old
    row.cells[1].text = new
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, '微软雅黑', 10)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('我们不是来"替换"酒店现有系统的，而是叠加在现有系统之上，用AI能力补齐酒店在数字化运营上的短板。')
set_run_font(run, '微软雅黑', 11, bold=True)

doc.add_paragraph()

# ===== 试点期间做什么 =====
h = doc.add_heading('试点期间，我们为酒店做什么（免费）', 2)
for run in h.runs:
    set_run_font(run, '微软雅黑', 14, bold=True)

sections = [
    ('① 收益管理与自动调价', [
        ('AI收益管理系统', '实时监控携程、美团、飞猪等平台上竞品的价格、挂牌等级、评价变化，根据市场需求、天气、节假日前后，自动给出定价建议。酒店收益人员只需确认执行，不需要手工收集数据、不需要做表格。'),
        ('自动调价预警', '竞品价格发生显著变化时，自动推送预警，建议调价幅度，一目了然。'),
    ]),
    ('② OTA智能维护', [
        ('差评实时监控与预警', '携程、美团等平台出现差评后，30分钟内通过微信推送给店长。AI自动分析问题原因（服务/设施/卫生/噪音），给出归因。AI生成差评回复建议，人工审核后发出。'),
        ('好评引导', '离店后自动推送点评邀请，配合小额优惠券。转化率目标：点评邀请打开率 > 40%。'),
        ('携程挂牌等级维护', '监控挂牌指标达标情况，提前预警可能掉级风险，给出冲刺挂牌等级的行动建议。'),
    ]),
    ('③ 私域运营自动化', [
        ('离店客人微信社群沉淀', '客人离店后，自动引导加入酒店微信社群。不需要前台手动操作，不需要客人主动留微信。'),
        ('社群自动化运营', '入住前一天自动发送入住指南；离店当天推送点评邀请；离店7天后推送复购优惠唤醒沉睡客人；日常定时推送天气预报、本地旅游信息、促销活动。'),
        ('RFM会员分层运营', '根据入住频次、消费金额、互动行为，对会员进行分层（A/B/C/D），不同层级推送不同优惠策略，提升复购率。'),
    ]),
    ('④ C端：客人自然语言预订与服务', [
        ('微信AI客服', '客人通过微信问入住时间/能不能带宠物/有停车场吗——AI自动识别意图，秒级回复。不需要打电线，不需要等前台回复。需要转人工时，一键转接。'),
        ('自然语言预订助手', '客人说"周六一家三口想住江景房，预算600左右"——AI自动理解需求，查询可用房态，推荐最合适的TOP3方案，客人确认后协助完成预订（跳转小程序支付）。全程在微信里完成。'),
        ('入住前/后自动服务', '入住前一天自动发送入住指南；入住当天推送房间号、Wi-Fi密码等；离店后自动推送点评邀请和复购优惠。'),
    ]),
    ('⑤ B端：管理者自然语言自动运营', [
        ('运营数据查询', '管理者问"本周入住率多少？""本月收益比上月多了还是少了？"AI直接给出答案，不需要找IT查报表。'),
        ('竞品动态追踪', '问"本周周边竞品价格有什么变化？"AI自动汇总携程、美团数据，生成竞品周报推送给管理者。'),
        ('自动生成月报', '每月自动生成核心指标报告：入住率/RevPAR/好评率/差评原因分布/竞品对比。管理者对运营现状有清晰数据支撑，不需要手工做表。'),
    ]),
]

for sec_title, items in sections:
    p = doc.add_paragraph()
    run = p.add_run(sec_title)
    set_run_font(run, '微软雅黑', 12, bold=True, color=RGBColor(0xC0, 0x00, 0x00))
    for t2, d in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(t2 + '：')
        set_run_font(run, '微软雅黑', 10, bold=True)
        run = p.add_run(d)
        set_run_font(run, '微软雅黑', 10)

doc.add_paragraph()

# ===== 酒店需要做什么 =====
h = doc.add_heading('酒店需要做什么（几乎没有额外工作）', 2)
for run in h.runs:
    set_run_font(run, '微软雅黑', 14, bold=True)

tbl3 = doc.add_table(rows=7, cols=3)
tbl3.style = 'Table Grid'
header_row(tbl3, ['责任方', '事项', '工作量'])

tasks = [
    ('我们承担', 'AI系统部署、技术维护、数据监控', '我们来做'),
    ('我们承担', '员工培训（如何使用AI工具）', '1次 x 2小时'),
    ('我们承担', '月度数据报告生成', '我们来做'),
    ('酒店配合', '确认一名对接人（运营或前厅负责人）', '日常沟通'),
    ('酒店配合', '初期提供基础房型/价格/设施资料', '约1小时'),
    ('酒店配合', '试运行期间反馈使用体验', '按需'),
]
for i, (who, what, work) in enumerate(tasks):
    row = tbl3.rows[i + 1]
    row.cells[0].text = who
    row.cells[1].text = what
    row.cells[2].text = work
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, '微软雅黑', 10)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('不需要：更换PMS系统 / 停用现有OTA渠道 / 购买硬件设备 / 一次性付费')
set_run_font(run, '微软雅黑', 11, bold=True, color=RGBColor(0xC0, 0x00, 0x00))

doc.add_paragraph()

# ===== 合作方式 =====
h = doc.add_heading('试点结束后的合作方式（双方自愿）', 2)
for run in h.runs:
    set_run_font(run, '微软雅黑', 14, bold=True)

levels = [
    ('基础版：', '收益管理 + OTA监控 + 月度报告', '[待定]/月'),
    ('标准版：', '基础版 + 私域运营 + C端AI预订客服 + B端AI运营助手', '[待定]/月'),
    ('定制版：', '根据酒店需求定制开发', '单独定价'),
]
for t2, content, price in levels:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(t2)
    set_run_font(run, '微软雅黑', 11, bold=True)
    run = p.add_run(content)
    set_run_font(run, '微软雅黑', 11)
    run = p.add_run('  ' + price)
    set_run_font(run, '微软雅黑', 11, bold=True, color=RGBColor(0xC0, 0x00, 0x00))

p = doc.add_paragraph()
run = p.add_run('我们希望用试点效果说话，不做强制推销。如果试点结束后酒店选择不续约，没有任何违约金或遗留问题。')
set_run_font(run, '微软雅黑', 11, bold=True)

doc.add_paragraph()

# ===== 为什么免费 =====
h = doc.add_heading('我们为什么愿意免费做试点', 2)
for run in h.runs:
    set_run_font(run, '微软雅黑', 14, bold=True)

reasons = [
    ('网络效应：', 'AHL的核心价值在于网络效应——接入的酒店越多，数据越丰富，AI越聪明。早期每一个试点酒店都是宝贵伙伴。'),
    ('信任建立：', '酒店行业的信任建立周期很长。用真实数据、真实效果，比任何PPT演示都有说服力。我们需要成功案例。'),
    ('产品迭代：', '试点本身是产品迭代的过程——酒店反馈的每一个问题，都是产品改进的方向。'),
]
for t2, d in reasons:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(t2)
    set_run_font(run, '微软雅黑', 11, bold=True)
    run = p.add_run(d)
    set_run_font(run, '微软雅黑', 11)

doc.add_paragraph()

# ===== 适合什么样的酒店 =====
h = doc.add_heading('什么样的酒店适合成为首批试点', 2)
for run in h.runs:
    set_run_font(run, '微软雅黑', 14, bold=True)

criteria = [
    '在本地市场有一定的客源基础（携程评分4.5以上更佳）',
    '管理团队对数字化运营有开放心态',
    '日常运营中有使用微信与客人沟通的场景',
    '有OTA在售房间（携程/美团/飞猪等平台）',
    '单体酒店或小型连锁（非大型集团总部统一管理项目）',
]
for item in criteria:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run('✅ ' + item)
    set_run_font(run, '微软雅黑', 11)

doc.add_paragraph()

# ===== 下一步 =====
h = doc.add_heading('接下来怎么开始', 2)
for run in h.runs:
    set_run_font(run, '微软雅黑', 14, bold=True)

steps = [
    ('第一步 互相了解（30分钟）', '我们了解酒店运营现状、痛点、期待；酒店了解我们的服务内容和合作方式'),
    ('第二步 确认范围（1天）', '共同确定试点周期、重点服务模块、评估指标，签订简单试点合作协议'),
    ('第三步 进场部署（1-3天）', '完成系统对接和初始化配置，对酒店相关人员进行培训，正式启动试点'),
    ('第四步 月度跟踪（持续）', '每月提供数据报告，每月一次复盘沟通，持续优化'),
]
for t2, d in steps:
    p = doc.add_paragraph()
    run = p.add_run(t2 + '：')
    set_run_font(run, '微软雅黑', 11, bold=True)
    run = p.add_run(d)
    set_run_font(run, '微软雅黑', 11)

doc.add_paragraph()

# ===== 核心优势总结 =====
h = doc.add_heading('核心优势总结', 2)
for run in h.runs:
    set_run_font(run, '微软雅黑', 14, bold=True)

tbl4 = doc.add_table(rows=6, cols=2)
tbl4.style = 'Table Grid'
header_row(tbl4, ['对酒店管理层', '对酒店业主方'])

adv = [
    ('收益管理：AI自动分析竞品、给出调价建议，一键执行', '收益可见：直接提升客房RevPAR，降低OTA依赖'),
    ('运营减负：AI客服7x24小时接待，减少人力消耗', '数据说话：每月核心指标报告，运营状况一目了然'),
    ('口碑改善：差评30分钟内响应，修复率提升', '竞争力强化：数字化能力领先周边竞品'),
    ('私域沉淀：离店客人自动入社群，复购触达不靠人工', '轻资产试水：免费试点，不投入任何前期成本'),
    ('自然语言交互：客人说话就能订房，管理者说话就能查数据', '长期绑定：试点转正式合作，建立长期数字化伙伴关系'),
]
for i, (m, o) in enumerate(adv):
    row = tbl4.rows[i + 1]
    row.cells[0].text = m
    row.cells[1].text = o
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, '微软雅黑', 10)

doc.add_paragraph()

# ===== 联系我们 =====
h = doc.add_heading('联系我们', 2)
for run in h.runs:
    set_run_font(run, '微软雅黑', 14, bold=True)

p = doc.add_paragraph()
run = p.add_run('如果贵酒店对成为首批试点合作伙伴感兴趣，欢迎联系我们了解详情。')
set_run_font(run, '微软雅黑', 11)

contacts = ['联系人：[联系人姓名]', '电话：17760348653（微信同号）', '邮箱：ericzhangshi@163.com']
for c in contacts:
    p = doc.add_paragraph()
    run = p.add_run(c)
    set_run_font(run, '微软雅黑', 11, bold=True)

p = doc.add_paragraph()
run = p.add_run('我们非常愿意到酒店拜访，当面沟通合作细节。')
set_run_font(run, '微软雅黑', 11, bold=True, color=RGBColor(0x2E, 0x75, 0xB6))

doc.add_paragraph()

# ===== 附：关于AHL =====
h = doc.add_heading('附：关于AHL项目', 2)
for run in h.runs:
    set_run_font(run, '微软雅黑', 12, bold=True)

p = doc.add_paragraph()
run = p.add_run(
    'AHL（AI Hotel Language）是一个去中心化旅行服务平台，致力于通过AI技术为酒店提供'
    '智能化的运营工具和私域流量运营能力。我们的团队成员有超过20年的酒店行业经验，'
    '深刻理解酒店的运营痛点和数字化需求。'
)
set_run_font(run, '微软雅黑', 10)

doc.add_paragraph()
f = doc.add_paragraph()
f.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = f.add_run('本合作意向书最终解释权归AHL项目团队所有。')
set_run_font(run, '微软雅黑', 9, color=RGBColor(0x80, 0x80, 0x80))

doc.save('docs/AHL-酒店智能升级-合作意向书-V1.0.docx')
print('OK')
