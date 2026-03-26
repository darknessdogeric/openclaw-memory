import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 标题
title = doc.add_heading('AHL商业模式深化方案 V2.0', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('基于酒店细分市场的完整计价体系')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('项目: AHL (AI Hotel Language)')
doc.add_paragraph('版本: V2.0 深化版')
doc.add_paragraph('日期: 2026年3月23日')
doc.add_paragraph('')

# 一、酒店细分市场定价矩阵
doc.add_heading('一、酒店细分市场定价矩阵', level=1)

doc.add_heading('1.1 按酒店档次分类（基于STR标准）', level=2)

doc.add_heading('奢华酒店（Luxury, ADR ¥1500+）', level=3)
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '服务模块'
hdr[1].text = '初装费'
hdr[2].text = '月度维护'
hdr[3].text = '月均TOKEN预算'
data = [
    ('基础版', '¥0', '¥2,000', '¥15,000-30,000'),
    ('进阶版', '¥0', '¥3,500', '¥30,000-60,000'),
    ('旗舰版', '¥0', '¥5,000', '¥60,000-100,000'),
    ('月均总成本', '-', '-', '¥17,000-105,000')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('高端酒店（Upper Upscale, ADR ¥800-1500）', level=3)
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '服务模块'
hdr[1].text = '初装费'
hdr[2].text = '月度维护'
hdr[3].text = '月均TOKEN预算'
data = [
    ('基础版', '¥0', '¥1,200', '¥8,000-15,000'),
    ('进阶版', '¥0', '¥2,000', '¥15,000-25,000'),
    ('旗舰版', '¥0', '¥3,000', '¥25,000-40,000'),
    ('月均总成本', '-', '-', '¥9,200-43,000')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('中高端酒店（Upscale, ADR ¥400-800）', level=3)
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '服务模块'
hdr[1].text = '初装费'
hdr[2].text = '月度维护'
hdr[3].text = '月均TOKEN预算'
data = [
    ('基础版', '¥0', '¥800', '¥5,000-10,000'),
    ('进阶版', '¥0', '¥1,200', '¥10,000-18,000'),
    ('旗舰版', '¥0', '¥1,800', '¥18,000-25,000'),
    ('月均总成本', '-', '-', '¥5,800-26,800')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('中端酒店（Upper Midscale, ADR ¥250-400）', level=3)
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '服务模块'
hdr[1].text = '初装费'
hdr[2].text = '月度维护'
hdr[3].text = '月均TOKEN预算'
data = [
    ('基础版', '¥0', '¥500', '¥3,000-6,000'),
    ('进阶版', '¥0', '¥800', '¥6,000-10,000'),
    ('旗舰版', '¥0', '¥1,200', '¥10,000-15,000'),
    ('月均总成本', '-', '-', '¥3,500-16,200')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('经济型酒店（Midscale/Economy, ADR ¥100-250）', level=3)
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '服务模块'
hdr[1].text = '初装费'
hdr[2].text = '月度维护'
hdr[3].text = '月均TOKEN预算'
data = [
    ('轻量版', '¥0', '¥0', '¥1,000-3,000'),
    ('基础版', '¥0', '¥200', '¥3,000-5,000'),
    ('标准版', '¥0', '¥400', '¥5,000-8,000'),
    ('月均总成本', '-', '-', '¥1,000-8,400')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

# 1.2 按物业类型分类
doc.add_heading('1.2 按物业类型分类', level=2)

doc.add_heading('商务酒店', level=3)
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '规模'
hdr[1].text = '房间数'
hdr[2].text = '推荐方案'
hdr[3].text = '月度成本区间'
data = [
    ('小型商务', '50-100间', '方案B基础版', '¥4,000-8,000'),
    ('中型商务', '100-200间', '方案B进阶版', '¥8,000-15,000'),
    ('大型商务', '200-400间', '方案C旗舰版', '¥15,000-30,000'),
    ('会议型', '400+间', '方案C全功能', '¥30,000-50,000')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('度假酒店', level=3)
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '类型'
hdr[1].text = '定位'
hdr[2].text = '推荐方案'
hdr[3].text = '月度成本区间'
data = [
    ('滨海度假', '亲子/海滩', '方案C旗舰版', '¥20,000-40,000'),
    ('山地度假', '滑雪/温泉', '方案B进阶版', '¥15,000-30,000'),
    ('乡村民宿', '田园/景区', '方案B基础版', '¥5,000-12,000'),
    ('城市周边', '周末/团建', '方案B基础版', '¥8,000-15,000')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

# 1.3 综合定价矩阵
doc.add_heading('1.3 综合定价矩阵表', level=2)
table = doc.add_table(rows=11, cols=6)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '酒店类型'
hdr[1].text = '规模'
hdr[2].text = '方案'
hdr[3].text = '月度维护'
hdr[4].text = 'TOKEN月费'
hdr[5].text = '月均总成本'
data = [
    ('奢华酒店', '100-300间', 'C-旗舰', '¥5,000', '¥60,000-100,000', '¥65,000-105,000'),
    ('高端酒店', '200-400间', 'C-进阶', '¥3,000', '¥25,000-40,000', '¥28,000-43,000'),
    ('中高端酒店', '100-200间', 'B-进阶', '¥1,200', '¥10,000-18,000', '¥11,200-19,200'),
    ('中端酒店', '80-150间', 'B-基础', '¥800', '¥6,000-10,000', '¥6,800-10,800'),
    ('经济型酒店', '50-100间', 'C-轻量', '¥200', '¥3,000-5,000', '¥3,200-5,200'),
    ('精品民宿', '10-30间', 'B-基础', '¥800', '¥3,000-6,000', '¥3,800-6,800'),
    ('景区民宿', '5-20间', 'C-轻量', '¥0', '¥1,500-4,000', '¥1,500-4,000'),
    ('度假酒店', '50-200间', 'C-进阶', '¥2,000', '¥15,000-30,000', '¥17,000-32,000'),
    ('公寓酒店', '50-150间', 'C-标准', '¥1,000', '¥8,000-15,000', '¥9,000-16,000'),
    ('会议酒店', '300+间', 'C-旗舰', '¥4,000', '¥30,000-50,000', '¥34,000-54,000')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

# 二、模块化定价
doc.add_heading('二、AGENT/SKILL模块化定价', level=1)

doc.add_heading('2.1 C端AI管家模块', level=2)
table = doc.add_table(rows=5, cols=5)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '模块'
hdr[1].text = '功能'
hdr[2].text = '月TOKEN消耗'
hdr[3].text = '适合档次'
hdr[4].text = '建议月费'
data = [
    ('客房服务', '预订/入住/退房', '300K-500K', '全档次', '¥1,500-4,000'),
    ('餐饮服务', '餐厅顾问/推荐', '200K-400K', '中高端+', '¥1,000-3,200'),
    ('礼宾服务', '本地向导/交通', '150K-300K', '中高端+', '¥750-2,400'),
    ('第四空间', '场景零售/推荐', '100K-200K', '中高端+', '¥500-1,600')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('2.2 B端AI运营官模块', level=2)
table = doc.add_table(rows=9, cols=5)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '模块'
hdr[1].text = '功能'
hdr[2].text = '降本增效价值'
hdr[3].text = '适合档次'
hdr[4].text = '建议月费'
data = [
    ('收益管理', '动态定价+库存', 'RevPAR+10-15%', '中端+', '¥2,000-5,000'),
    ('营销获客', '内容生成+OTA', '获客成本-20%', '全档次', '¥1,500-4,000'),
    ('客服响应', '7×24智能客服', '人工-50%', '全档次', '¥1,000-3,000'),
    ('财务管理', '自动对账+报表', '效率+40%', '中端+', '¥800-2,500'),
    ('工程运维', '设备监控+能耗', '能耗-15%', '中大型', '¥1,200-3,000'),
    ('人资排班', '智能排班+招聘', '人效+20%', '全档次', '¥800-2,000'),
    ('宴会销售', '婚宴/会议统筹', '收入+25%', '度假/会议', '¥1,500-4,000'),
    ('空间运营', '会议室/闲置', '坪效+10%', '全档次', '¥600-1,500')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('2.3 模块组合定价', level=2)
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '套餐类型'
hdr[1].text = '包含模块'
hdr[2].text = '月费+TOKEN'
hdr[3].text = '月均总成本'
data = [
    ('基础套餐', '客房+基础客服', '¥300-500+¥1,000-2,000', '¥1,300-2,500'),
    ('标准套餐', '客房+餐饮+收益+营销', '¥1,500+¥5,000-10,000', '¥6,500-11,500'),
    ('进阶套餐', '全C端+核心B端', '¥2,500+¥12,000-20,000', '¥14,500-22,500'),
    ('旗舰套餐', '全模块+定制', '¥5,000+¥30,000-100,000', '¥35,000-105,000')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

# 三、发展阶段定价
doc.add_heading('三、发展阶段定价策略', level=1)

doc.add_heading('3.1 种子期（M1-M3）', level=2)
doc.add_paragraph('目标: 50家商户入驻，验证产品价值')
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'
data = [
    ('策略', '全部免费或极低价'),
    ('收入预期', '接近零，重在市场验证')
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_heading('3.2 验证期（M4-M6）', level=2)
doc.add_paragraph('目标: 200家商户，月GMV验证')
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'
data = [
    ('策略', '低价切入，月费¥300-500 + TOKEN¥0.003-0.004'),
    ('收入预期', '¥50-100万/月')
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_heading('3.3 增长期（M7-M18）', level=2)
doc.add_paragraph('目标: 2000+商户，区域领先')
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
data = [
    ('M7-M12', '方案B：月费¥500-1500 + TOKEN¥0.004-0.005'),
    ('M13-M18', '方案B→C过渡：月费¥200-800 + TOKEN¥0.005-0.006'),
    ('收入预期', 'M12: ¥500-800万/月, M18: ¥2,000-3,000万/月')
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_heading('3.4 成熟期（M19-M36）', level=2)
doc.add_paragraph('目标: 5000+商户，实现盈利')
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
data = [
    ('M19-M24', '方案C为主：纯TOKEN¥0.006-0.007'),
    ('M25-M36', '方案C：TOKEN¥0.007-0.008 + 增值服务'),
    ('收入预期', 'M24: ¥4,000-6,000万/月, M36: ¥6,500万+/月')
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

# 四、方案对比
doc.add_heading('四、三方案完整对比', level=1)

doc.add_heading('4.1 核心指标对比', level=2)
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '指标'
hdr[1].text = '方案A'
hdr[2].text = '方案B'
hdr[3].text = '方案C'
data = [
    ('36个月累计', '¥2.85亿', '¥5.8亿', '¥11.2亿'),
    ('月均收入', '¥79万', '¥161万', '¥311万'),
    ('毛利率', '52%', '16%', '46%')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('4.2 酒店类型适配度', level=2)
table = doc.add_table(rows=10, cols=5)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '酒店类型'
hdr[1].text = '方案A'
hdr[2].text = '方案B'
hdr[3].text = '方案C'
hdr[4].text = '推荐'
data = [
    ('奢华酒店', '⭐', '⭐⭐⭐', '⭐⭐⭐⭐⭐', 'C'),
    ('高端酒店', '⭐⭐', '⭐⭐⭐⭐', '⭐⭐⭐⭐', 'B/C'),
    ('中高端酒店', '⭐⭐⭐', '⭐⭐⭐⭐⭐', '⭐⭐⭐⭐', 'B'),
    ('中端酒店', '⭐⭐⭐⭐', '⭐⭐⭐⭐⭐', '⭐⭐⭐', 'B'),
    ('经济型', '⭐⭐⭐⭐', '⭐⭐⭐', '⭐⭐⭐⭐⭐', 'C'),
    ('民宿', '⭐⭐⭐', '⭐⭐⭐⭐⭐', '⭐⭐⭐⭐', 'B/C'),
    ('度假酒店', '⭐⭐', '⭐⭐⭐⭐', '⭐⭐⭐⭐⭐', 'C'),
    ('公寓酒店', '⭐⭐', '⭐⭐⭐⭐', '⭐⭐⭐⭐⭐', 'C'),
    ('会议酒店', '⭐⭐', '⭐⭐⭐⭐', '⭐⭐⭐⭐⭐', 'C')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('4.3 投资回报分析', level=2)
table = doc.add_table(rows=4, cols=5)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '酒店类型'
hdr[1].text = '月均成本'
hdr[2].text = '月均收益'
hdr[3].text = 'ROI'
hdr[4].text = '回本周期'
data = [
    ('中端100间', '¥8,800', '¥7,000-12,000', '80%-136%', '1-1.4月'),
    ('中高端150间', '¥15,000', '¥17,500-28,000', '117%-187%', '0.5-0.9月'),
    ('精品民宿20间', '¥2,500', '¥3,000-5,000', '120%-200%', '0.5-0.8月')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

# 五、最终建议
doc.add_heading('五、最终建议', level=1)

doc.add_heading('5.1 最佳路径', level=2)
doc.add_paragraph('阶段一（M1-M6）: 全部免费 - 快速获取种子用户')
doc.add_paragraph('阶段二（M7-M12）: 方案B - 月费¥500-1500 + TOKEN¥0.004')
doc.add_paragraph('阶段三（M13-M24）: 方案B→C - 月费¥200-800 + TOKEN¥0.005')
doc.add_paragraph('阶段四（M25+）: 方案C为主 - 纯TOKEN¥0.006-0.007 + 增值服务')

doc.add_heading('5.2 核心洞察', level=2)
doc.add_paragraph('1. TOKEN是趋势 - AI使用不可逆，TOKEN消耗会指数增长')
doc.add_paragraph('2. 酒店类型决定定价 - 档次越高、规模越大，TOKEN消耗越高')
doc.add_paragraph('3. 规模效应明显 - 商户越多，TOKEN成本越低，毛利越高')
doc.add_paragraph('4. 投资回报快速 - 大多数酒店1-2个月即可回本')

# 保存
desktop = os.path.join(os.environ['USERPROFILE'], 'Desktop')
output_path = os.path.join(desktop, 'AHL-商业模式深化方案V2.0.docx')
doc.save(output_path)
print(f'已保存到: {output_path}')
