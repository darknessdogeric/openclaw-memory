import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 标题
title = doc.add_heading('AHL商务合作收费标准 V3.0', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('完整版：TOKEN计费+三种模式对比+分级体系')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('项目: AHL (AI Hotel Language)')
doc.add_paragraph('版本: V3.0 正式版')
doc.add_paragraph('日期: 2026年3月23日')
doc.add_paragraph('')

# 一、底层逻辑分析
doc.add_heading('一、底层逻辑分析', level=1)

doc.add_heading('1.1 收费原理', level=2)
doc.add_paragraph('为什么收费？')
doc.add_paragraph('- AHL提供的是AI能力，本质是大模型TOKEN消耗')
doc.add_paragraph('- TOKEN是AI运行的"燃料"，按量收费是行业趋势')
doc.add_paragraph('- 酒店使用AI越多，价值创造越大，收费越合理')

doc.add_heading('1.2 月度维护费 vs TOKEN费', level=2)
table = doc.add_table(rows=3, cols=4)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '费用类型'
hdr[1].text = '定价依据'
hdr[2].text = '收费逻辑'
hdr[3].text = '合理性'
data = [
    ('月度维护费', '房间数量', '系统运维、客服成本', '与房间数正相关'),
    ('TOKEN费', '房间+档次', '大模型调用量、AI复杂度', '使用越多付费越多')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

# 二、月度维护费标准
doc.add_heading('二、月度维护费标准（按房间数量）', level=1)

doc.add_heading('2.1 定价原则', level=2)
doc.add_paragraph('为什么按房间数量？')
doc.add_paragraph('1. 房间数量直接决定系统负载')
doc.add_paragraph('2. 房间数量是酒店行业通用指标')
doc.add_paragraph('3. 与成本对应，公平合理')

doc.add_heading('2.2 收费标准', level=2)
table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '房间数量区间'
hdr[1].text = '月度维护费'
hdr[2].text = '说明'
data = [
    ('50间以下', '¥300/月', '轻量版'),
    ('50-100间', '¥500/月', '基础版'),
    ('100-150间', '¥800/月', '标准版'),
    ('150-200间', '¥1,200/月', '进阶版'),
    ('200-300间', '¥1,800/月', '商务版'),
    ('300-500间', '¥2,800/月', '旗舰版'),
    ('500间以上', '¥5,000/月', '至尊版')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

# 三、TOKEN计费标准
doc.add_heading('三、TOKEN计费标准', level=1)

doc.add_heading('3.1 定价原则', level=2)
doc.add_paragraph('为什么按"房间数量 + 酒店档次"？')
doc.add_paragraph('1. 房间数量决定基础调用量')
doc.add_paragraph('2. 酒店档次决定AI服务复杂度')
doc.add_paragraph('3. 不同档次支付意愿和能力不同')

doc.add_heading('3.2 TOKEN消耗估算', level=2)
table = doc.add_table(rows=7, cols=5)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '酒店档次'
hdr[1].text = 'ADR区间'
hdr[2].text = '月均TOKEN'
hdr[3].text = '单价'
hdr[4].text = '月费区间'
data = [
    ('奢华', '¥1500+', '150K-300K', '¥0.008', '¥1,200-2,400'),
    ('高端', '¥800-1500', '80K-180K', '¥0.007', '¥560-1,260'),
    ('中高端', '¥400-800', '50K-120K', '¥0.006', '¥300-720'),
    ('中端', '¥250-400', '30K-70K', '¥0.005', '¥150-350'),
    ('经济型', '¥100-250', '15K-40K', '¥0.004', '¥60-160'),
    ('民宿', '差异大', '20K-60K', '¥0.004', '¥80-240')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

# 四、三种商业模式对比
doc.add_heading('四、三种商业模式对比', level=1)

doc.add_heading('4.1 方案A：初装费 + 月度维护费（传统SaaS模式）', level=2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
data = [
    ('项目', '定价'),
    ('初装费', '¥15,000-50,000'),
    ('月度维护费', '¥500-5,000（按房间数）'),
    ('TOKEN费', '包含在月费中，限量'),
    ('适用场景', '预算充足的传统酒店')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

doc.add_heading('4.2 方案B：免初装 + 月度维护 + TOKEN计费（混合模式）⭐推荐', level=2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
data = [
    ('项目', '定价'),
    ('初装费', '¥0（免费部署）'),
    ('月度维护费', '¥300-5,000（按房间数）'),
    ('TOKEN计费', '¥0.004-0.008（按档次）'),
    ('适用场景', '大多数中小型酒店')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

doc.add_heading('4.3 方案C：免初装 + 免维护 + TOKEN计费（AI原生模式）', level=2)
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
data = [
    ('项目', '定价'),
    ('初装费', '¥0（完全免费）'),
    ('月度维护费', '¥0（不收取）'),
    ('TOKEN计费', '¥0.006-0.010（按档次）'),
    ('适用场景', '追求极致低门槛的商户')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i].cells[j].text = val

# 五、综合定价矩阵
doc.add_heading('五、综合定价矩阵', level=1)

doc.add_heading('5.1 完整定价表', level=2)
table = doc.add_table(rows=18, cols=5)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '酒店档次'
hdr[1].text = '房间数'
hdr[2].text = '月度维护'
hdr[3].text = 'TOKEN月费'
hdr[4].text = '月均总费用'
data = [
    ('奢华', '100-200', '¥1,200', '¥1,200-2,000', '¥2,400-3,200'),
    ('奢华', '200-300', '¥1,800', '¥2,000-3,000', '¥3,800-4,800'),
    ('奢华', '300+', '¥2,800', '¥3,000-5,000', '¥5,800-7,800'),
    ('高端', '100-150', '¥800', '¥800-1,200', '¥1,600-2,000'),
    ('高端', '150-200', '¥1,200', '¥1,000-1,500', '¥2,200-2,700'),
    ('高端', '200-400', '¥1,800', '¥1,500-2,500', '¥3,300-4,300'),
    ('中高端', '80-100', '¥500', '¥400-700', '¥900-1,200'),
    ('中高端', '100-150', '¥800', '¥500-900', '¥1,300-1,700'),
    ('中高端', '150-200', '¥1,200', '¥700-1,200', '¥1,900-2,400'),
    ('中端', '50-80', '¥300', '¥200-350', '¥500-650'),
    ('中端', '80-100', '¥500', '¥250-450', '¥750-950'),
    ('中端', '100-150', '¥800', '¥350-600', '¥1,150-1,400'),
    ('经济型', '50以下', '¥300', '¥80-150', '¥380-450'),
    ('经济型', '50-100', '¥500', '¥100-250', '¥600-750'),
    ('民宿', '10-20', '¥300', '¥100-200', '¥400-500'),
    ('民宿', '20-30', '¥500', '¥150-300', '¥650-800'),
    ('民宿', '30+', '¥800', '¥200-400', '¥1,000-1,200')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('5.2 投资回报估算', level=2)
table = doc.add_table(rows=7, cols=5)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '酒店类型'
hdr[1].text = '月均成本'
hdr[2].text = '预期收益提升'
hdr[3].text = 'ROI'
hdr[4].text = '回本周期'
data = [
    ('奢华酒店(200间)', '¥4,300', '¥15,000-30,000', '250%-600%', '0.1-0.3月'),
    ('高端酒店(200间)', '¥3,300', '¥12,000-20,000', '260%-500%', '0.2-0.3月'),
    ('中高端酒店(150间)', '¥1,700', '¥8,000-15,000', '370%-780%', '0.1-0.2月'),
    ('中端酒店(100间)', '¥950', '¥5,000-10,000', '426%-953%', '0.1-0.2月'),
    ('经济型(80间)', '¥600', '¥3,000-6,000', '400%-900%', '0.1-0.2月'),
    ('民宿(20间)', '¥500', '¥2,000-5,000', '300%-900%', '0.1-0.3月')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

# 六、阶段定价策略
doc.add_heading('六、阶段定价策略', level=1)

doc.add_heading('6.1 种子期（M1-M3）', level=2)
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'
data = [
    ('策略', '全部免费'),
    ('目标', '50家标杆')
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_heading('6.2 验证期（M4-M6）', level=2)
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'
data = [
    ('策略', '月费5折 + TOKEN5折'),
    ('目标', '200家')
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_heading('6.3 增长期（M7-M18）', level=2)
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
data = [
    ('M7-M12', '月费8折 + TOKEN8折'),
    ('M13-M18', '月费9折 + TOKEN9折'),
    ('目标', '2000家')
]
for i, (k, v) in enumerate(data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_heading('6.4 成熟期（M19+）', level=2)
doc.add_paragraph('全价执行，按标准定价')

# 七、总结
doc.add_heading('七、总结', level=1)

doc.add_heading('7.1 定价原则', level=2)
table = doc.add_table(rows=3, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '费用类型'
hdr[1].text = '定价依据'
hdr[2].text = '底层逻辑'
data = [
    ('月度维护费', '房间数量', '系统运维成本与房间数正相关'),
    ('TOKEN费', '房间数量+酒店档次', 'AI调用量与使用深度正相关')
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('7.2 推荐方案', level=2)
doc.add_paragraph('首选方案B（混合模式）')
doc.add_paragraph('- 初装免费降低门槛')
doc.add_paragraph('- 月度维护费覆盖基础成本')
doc.add_paragraph('- TOKEN费体现价值创造')

# 保存
desktop = os.path.join(os.environ['USERPROFILE'], 'Desktop')
output_path = os.path.join(desktop, 'AHL-商务合作收费标准V3.0.docx')
doc.save(output_path)
print(f'已保存到: {output_path}')
