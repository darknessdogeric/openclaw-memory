from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 创建文档
doc = Document()

# 标题
title = doc.add_heading('AHL 产品清单', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 副标题
subtitle = doc.add_paragraph('Product Catalog V1.0')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(12)
subtitle.runs[0].font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph('更新日期: 2026-03-17')
doc.add_paragraph('定位: AHL可插拔产品矩阵，支持模块化组合与快速部署')
doc.add_paragraph('架构: 大模型中枢 + 双AGENT + 细分场景SKILL')

doc.add_page_break()

# 目录
doc.add_heading('目录', 1)
doc.add_paragraph('一、平台核心产品 (Platform Core)')
doc.add_paragraph('二、C端产品矩阵 (Consumer Products)')
doc.add_paragraph('三、B端产品矩阵 (Business Products)')
doc.add_paragraph('四、行业解决方案 (Industry Solutions)')
doc.add_paragraph('五、数据产品 (Data Products)')
doc.add_paragraph('六、基础设施产品 (Infrastructure)')

doc.add_page_break()

# 一、平台核心产品
doc.add_heading('一、平台核心产品 (Platform Core)', 1)

doc.add_heading('1.1 双边匹配引擎 (Matching Engine)', 2)
table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '产品模块'
hdr_cells[1].text = '功能说明'
hdr_cells[2].text = '技术特性'
hdr_cells[3].text = '适用场景'

rows = [
    ['AHL-Match-001\n事实向量匹配引擎', '基于C端需求向量与B端供给向量的精准匹配计算', '多空间向量索引、实时相似度计算、动态权重调整', '全平台酒店搜索与推荐'],
    ['AHL-Match-002\n场景适配引擎', '7大出行场景的智能识别与匹配', '场景强度评分、多维度对齐算法', '场景化酒店筛选与排序'],
    ['AHL-Match-003\n实时动态匹配', '基于实时库存、价格、服务能力的动态匹配优化', '流式数据接入、毫秒级响应、供需平衡算法', '闪购/尾房/即时预订场景'],
    ['AHL-Match-004\n长尾匹配引擎', '小众需求与小众供给的精准对接', '稀疏数据建模、冷启动策略', '特色民宿/主题酒店/非标住宿']
]

for i, row_data in enumerate(rows, 1):
    cells = table.rows[i].cells
    for j, text in enumerate(row_data):
        cells[j].text = text

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('产品组合建议:').bold = True
doc.add_paragraph('• 标准版: Match-001 + Match-002')
doc.add_paragraph('• 进阶版: Match-001 + Match-002 + Match-003')
doc.add_paragraph('• 完整版: 全部4个模块')

# 1.2 C端AI管家
doc.add_heading('1.2 C端AI管家 (Consumer AI Butler)', 2)
table = doc.add_table(rows=7, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '产品模块'
hdr_cells[1].text = '功能说明'
hdr_cells[2].text = '核心能力'
hdr_cells[3].text = '输出形态'

rows = [
    ['AHL-C-001\n智能行程规划', '基于用户意图的全程行程规划', '意图识别、多目标优化、实时信息整合', '对话交互 + 可视化行程单'],
    ['AHL-C-002\n酒店智能推荐', '多轮对话理解需求，精准推荐匹配酒店', '对话状态管理、需求澄清、个性化排序', '推荐列表 + 推荐理由'],
    ['AHL-C-003\n预订全流程助手', '从搜索到下单的全程陪伴', '流程编排、异常处理、支付辅助', '嵌入式助手 + 主动提醒'],
    ['AHL-C-004\n入住前服务', '入住前48小时至入住当天的服务准备', '智能Check-in、需求预采集、服务预约', '消息推送 + 小程序/H5'],
    ['AHL-C-005\n入住中服务', '客房服务、本地向导、问题处理的实时响应', '客房控制集成、本地知识库、紧急支援', '语音/文字对话 + 服务调度'],
    ['AHL-C-006\n离店与复购', '退房办理、发票申请、下次出行提醒', '自动化退房、会员权益、复购预测', '消息推送 + 优惠券发放']
]

for i, row_data in enumerate(rows, 1):
    cells = table.rows[i].cells
    for j, text in enumerate(row_data):
        cells[j].text = text

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('产品组合建议:').bold = True
doc.add_paragraph('• 轻量版: C-002 + C-003')
doc.add_paragraph('• 标准版: C-001 + C-002 + C-003 + C-006')
doc.add_paragraph('• 完整版: 全部6个模块')

# 1.3 B端AI运营官
doc.add_heading('1.3 B端AI运营官 (Business AI Officer)', 2)
table = doc.add_table(rows=7, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '产品模块'
hdr_cells[1].text = '功能说明'
hdr_cells[2].text = '核心能力'
hdr_cells[3].text = '输出形态'

rows = [
    ['AHL-B-001\n智能定价中枢', '基于供需预测、竞争情报的动态定价建议', '需求预测、价格弹性建模、竞品监控', '定价建议 + 自动调价'],
    ['AHL-B-002\n库存优化引擎', '多渠道库存分配、房型升级策略、超售管理', '渠道ROI分析、库存释放策略、风险评估', '库存分配方案 + 预警'],
    ['AHL-B-003\n运营诊断助手', '日常运营数据监控、异常检测、改进建议', 'KPI监控、根因分析、最佳实践推荐', '日报/周报 + 改进清单'],
    ['AHL-B-004\n客户服务中枢', '客户咨询自动响应、投诉处理、满意度管理', 'NLP理解、情感分析、工单流转', '自动回复 + 人工升级'],
    ['AHL-B-005\n营销自动化', '促销活动策划、内容生成、渠道投放优化', '活动ROI预测、内容AIGC、投放优化', '活动方案 + 执行监控'],
    ['AHL-B-006\n数据报表中心', '自定义报表、数据可视化、多维度分析', '拖拽式报表、实时数据、下钻分析', '仪表盘 + 导出报告']
]

for i, row_data in enumerate(rows, 1):
    cells = table.rows[i].cells
    for j, text in enumerate(row_data):
        cells[j].text = text

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('产品组合建议:').bold = True
doc.add_paragraph('• 基础版: B-003 + B-006')
doc.add_paragraph('• 标准版: B-001 + B-002 + B-003 + B-006')
doc.add_paragraph('• 完整版: 全部6个模块')

# 1.4 数据底座
doc.add_heading('1.4 数据底座 (Data Foundation)', 2)
table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '产品模块'
hdr_cells[1].text = '功能说明'
hdr_cells[2].text = '技术特性'
hdr_cells[3].text = '数据规模'

rows = [
    ['AHL-D-001\n向量数据库', '酒店与用户的多维向量存储与检索', '高维向量索引、近似最近邻搜索、实时更新', '支持亿级向量'],
    ['AHL-D-002\n事实数据湖', '酒店基础信息、交易数据、行为数据的统一存储', '多源接入、Schema管理、数据治理', 'PB级存储'],
    ['AHL-D-003\n标签与画像系统', '酒店标签体系、用户画像的生成与管理', '自动标签提取、画像更新、标签推理', '千级标签维度'],
    ['AHL-D-004\n实时数据流', '库存、价格、订单的实时数据接入与处理', '流式计算、低延迟、Exactly-Once', '毫秒级延迟'],
    ['AHL-D-005\n数据飞轮引擎', '双边交互数据的反馈闭环与模型优化', 'A/B测试、模型迭代、效果归因', '持续自优化']
]

for i, row_data in enumerate(rows, 1):
    cells = table.rows[i].cells
    for j, text in enumerate(row_data):
        cells[j].text = text

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('产品组合建议:').bold = True
doc.add_paragraph('• 基础版: D-001 + D-002')
doc.add_paragraph('• 标准版: D-001 + D-002 + D-003 + D-004')
doc.add_paragraph('• 完整版: 全部5个模块')

doc.add_page_break()

# 二、C端产品矩阵
doc.add_heading('二、C端产品矩阵 (Consumer Products)', 1)

doc.add_heading('2.1 智能预订产品 (Smart Booking)', 2)
table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '产品代码'
hdr_cells[1].text = '产品名称'
hdr_cells[2].text = '功能说明'
hdr_cells[3].text = '定价模式'

rows = [
    ['AHL-CB-001', '智能搜索', '自然语言搜索、语义理解、意图识别', '免费/按调用量'],
    ['AHL-CB-002', '智能筛选', '多维度筛选、场景化筛选、智能排序', '免费'],
    ['AHL-CB-003', '比价助手', '多平台价格对比、历史价格趋势、最佳预订时机', '免费/增值服务'],
    ['AHL-CB-004', '闪购预订', '尾房闪购、限时特惠、库存清仓', '交易佣金'],
    ['AHL-CB-005', '企业预订', '企业协议价、差旅政策、统一结算', 'SaaS订阅']
]

for i, row_data in enumerate(rows, 1):
    cells = table.rows[i].cells
    for j, text in enumerate(row_data):
        cells[j].text = text

doc.add_heading('2.2 场景化服务产品 (Scenario Services)', 2)
table = doc.add_table(rows=8, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '产品代码'
hdr_cells[1].text = '产品名称'
hdr_cells[2].text = '功能说明'
hdr_cells[3].text = '目标场景'

rows = [
    ['AHL-CS-001', '商务出行包', '快速入住、会议室预订、商务中心、接送机', '商务差旅'],
    ['AHL-CS-002', '亲子度假包', '儿童设施、亲子活动、加床服务、周边景点', '家庭亲子'],
    ['AHL-CS-003', '情侣浪漫包', '景观房、浴缸、鲜花布置、浪漫晚餐', '情侣度假'],
    ['AHL-CS-004', '长者康养包', '无障碍设施、医疗协助、慢节奏服务、健康餐饮', '银发康养'],
    ['AHL-CS-005', '宠物友好包', '宠物政策、宠物设施、周边友好场所', '携宠出行'],
    ['AHL-CS-006', '长住旅居包', '周租/月租优惠、厨房设施、办公空间、社区活动', '数字游民/长住客'],
    ['AHL-CS-007', '社交体验包', '公区活动、社群匹配、当地体验、共享空间', '社交旅行']
]

for i, row_data in enumerate(rows, 1):
    cells = table.rows[i].cells
    for j, text in enumerate(row_data):
        cells[j].text = text

doc.add_heading('2.3 会员与忠诚度产品 (Membership & Loyalty)', 2)
table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '产品代码'
hdr_cells[1].text = '产品名称'
hdr_cells[2].text = '功能说明'
hdr_cells[3].text = '商业模式'

rows = [
    ['AHL-CL-001', 'AHL会员体系', '积分累积、等级权益、专属优惠', '平台会员'],
    ['AHL-CL-002', '酒店直销会员', '帮助酒店建立自有会员体系、私域运营', 'SaaS+交易佣金'],
    ['AHL-CL-003', '企业会员计划', '企业客户专属权益、协议价管理、消费分析', 'B2B订阅'],
    ['AHL-CL-004', '联合会员', '与航司/OTA/信用卡的会员互通、权益共享', '合作分成'],
    ['AHL-CL-005', '订阅制住宿', '月度/季度订阅、无限次入住、灵活退改', '订阅收入']
]

for i, row_data in enumerate(rows, 1):
    cells = table.rows[i].cells
    for j, text in enumerate(row_data):
        cells[j].text = text

doc.add_heading('2.4 内容发现产品 (Content Discovery)', 2)
table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '产品代码'
hdr_cells[1].text = '产品名称'
hdr_cells[2].text = '功能说明'
hdr_cells[3].text = '内容形式'

rows = [
    ['AHL-CC-001', '酒店故事', '酒店主理人故事、设计理念、在地文化', '图文/视频'],
    ['AHL-CC-002', '真实评价', '结构化评价、标签云、可信度评分', 'UGC'],
    ['AHL-CC-003', '探店内容', 'KOL探店、体验报告、避坑指南', '视频/直播'],
    ['AHL-CC-004', '目的地指南', '城市攻略、周边玩法、隐藏景点', '图文/地图'],
    ['AHL-CC-005', 'AI旅行日记', '自动生成旅行记录、照片整理、行程回顾', '智能生成']
]

for i, row_data in enumerate(rows, 1):
    cells = table.rows[i].cells
    for j, text in enumerate(row_data):
        cells[j].text = text

doc.add_page_break()

# 三、B端产品矩阵
doc.add_heading('三、B端产品矩阵 (Business Products)', 1)

doc.add_heading('3.1 智能运营产品 (Smart Operations)', 2)
table = doc.add_table(rows=10, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '产品代码'
hdr_cells[1].text = '产品名称'
hdr_cells[2].text = '功能说明'
hdr_cells[3].text = '适用部门'

rows = [
    ['AHL-BO-001', '前厅运营助手', '智能排房、夜班值守、客诉处理、VIP识别', '前厅部'],
    ['AHL-BO-002', '客房管理助手', '智能排班、客房检查、布草管理、维保触发', '客房部'],
    ['AHL-BO-003', '餐饮服务助手', '餐位管理、菜单工程、库存预警、宴会销售', '餐饮部'],
    ['AHL-BO-004', '宴会销售助手', '婚宴/寿宴/商务宴销售+LTV运营', '宴会部'],
    ['AHL-BO-005', '工程维保助手', '预测性维护、工单管理、能耗监控', '工程部'],
    ['AHL-BO-006', '安保监控助手', '消防管理、巡逻管理、应急预案、监控分析', '安保部'],
    ['AHL-BO-007', '质检培训助手', '客房质检、前厅质检、培训管理、OTA评价管理', '质检培训部'],
    ['AHL-BO-008', '行政后勤助手', '采购管理、资产管理、档案管理、印章管理', '行政部'],
    ['AHL-BO-009', '康乐康养助手', '康乐设施管理、康养服务、医养结合', '康乐部']
]

for i, row_data in enumerate(rows, 1):
    cells = table.rows[i].cells
    for j, text in enumerate(row_data):
        cells[j].text = text

doc.add_heading('3.2 收益管理产品 (Revenue Management)', 2)
table = doc.add_table(rows=7, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '产品代码'
hdr_cells[1].text = '产品名称'
hdr_cells[2].text = '功能说明'
hdr_cells[3].text = '核心算法'

rows = [
    ['AHL-BR-001', '动态定价引擎', '基于供需预测、竞争情报的实时定价建议', '需求预测模型、价格弹性模型'],
    ['AHL-BR-002', '库存优化系统', '多渠道库存分配、房型升级策略', '渠道ROI模型、库存释放算法'],
    ['AHL-BR-003', '竞争情报监控', '竞品价格监控、市场动态分析', '爬虫+NLP'],
    ['AHL-BR-004', '细分市场定价', '企业客户/会员/OTA差异化定价', '细分需求模型'],
    ['AHL-BR-005', '餐饮收益管理', '餐厅定价、套餐设计、时段优化', '餐饮需求预测'],
    ['AHL-BR-006', '预测与报告', '入住率预测、RevPAR预测、业绩归因', '时间序列模型']
]

for i, row_data in enumerate(rows, 1):
    cells = table.rows[i].cells
    for j, text in enumerate(row_data):
        cells[j].text = text

doc.add_heading('3.3 营销获客产品 (Marketing & Acquisition)', 2)
table = doc.add_table(rows=11, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '产品代码'
hdr_cells[1].text = '产品名称'
hdr_cells[2].text = '功能说明'
hdr_cells[3].text = '渠道覆盖'

rows = [
    ['AHL-BM-001', '获客雷达', '7×24小时监测潜在客户需求、竞品动态', '全渠道'],
    ['AHL-BM-002', '内容生成器', '营销文案、图片、视频的智能生成', '社媒/OTA'],
    ['AHL-BM-003', '社媒运营助手', '多平台账号管理、内容发布、互动管理', '微信/抖音/小红书'],
    ['AHL-BM-004', '会员运营助手', '会员生命周期管理、精准营销、流失预警', 'CRM'],
    ['AHL-BM-005', '活动运营助手', '促销活动策划、执行、效果分析', '全渠道'],
    ['AHL-BM-006', 'KOL合作管理', 'KOL筛选、合作管理、效果追踪', '达人营销'],
    ['AHL-BM-007', 'OTA运营助手', '携程/美团/飞猪等平台的运营优化', 'OTA'],
    ['AHL-BM-008', 'B2B市场助手', 'RFP响应、TMC对接、协议客户开发', 'B2B'],
    ['AHL-BM-009', '闲置空间运营', '会议室/包间/大堂吧等空间的高效利用', '空间租赁'],
    ['AHL-BM-010', '传讯品牌助手', '公关管理、品牌视觉、活动策划、新媒体', '品牌部']
]

for i, row_data in enumerate(rows, 1):
    cells = table.rows[i].cells
    for j, text in enumerate(row_data):
        cells[j].text = text

doc.add_heading('3.4 数据智能产品 (Data Intelligence)', 2)
table = doc.add_table(rows=6, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '产品代码'
hdr_cells[1].text = '产品名称'
hdr_cells[2].text = '功能说明'
hdr_cells[3].text = '数据输出'

rows = [
    ['AHL-BD-001', '经营分析仪表盘', '核心KPI监控、趋势分析、对标分析', '可视化报表'],
    ['AHL-BD-002', '客户画像分析', '客群分析、偏好洞察、价值分层', '画像报告'],
    ['AHL-BD-003', '市场洞察报告', '市场趋势、竞争格局、机会识别', '研究报告'],
    ['AHL-BD-004', '预测性分析', '需求预测、流失预测、异常预警', '预测模型'],
    ['AHL-BD-005', '自动化报表', '定制化报表、定时推送、多格式导出', '报表文件']
]

for i, row_data in enumerate(rows, 1):
    cells = table.rows[i].cells
    for j, text in enumerate(row_data):
        cells[j].text = text

doc.add_page_break()

# 保存文档
doc.save('C:/Users/Administrator/Desktop/AHL-Product-Catalog.docx')
print('Word文档已生成: C:/Users/Administrator/Desktop/AHL-Product-Catalog.docx')
