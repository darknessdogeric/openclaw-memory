#!/usr/bin/env python3
"""
NotebookLM - Intelligent Document Analysis and Synthesis Tool
类似Google NotebookLM的本地实现
"""

import os
import re
import json
import hashlib
from pathlib import Path
from typing import List, Dict, Optional, Tuple, Any
from dataclasses import dataclass, asdict
from datetime import datetime
import tempfile

# 文档处理
import PyPDF2
from docx import Document as DocxDocument

# 文本处理
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from collections import Counter

# 向量化和语义搜索
try:
    from sentence_transformers import SentenceTransformer
    from sklearn.metrics.pairwise import cosine_similarity
    EMBEDDINGS_AVAILABLE = True
except ImportError:
    EMBEDDINGS_AVAILABLE = False

# AI模型接口
try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False


@dataclass
class SourceDocument:
    """源文档"""
    id: str
    filename: str
    filepath: str
    content: str
    doc_type: str  # pdf, txt, md, docx
    metadata: Dict
    chunks: List[str]
    embeddings: Optional[List[List[float]]] = None
    created_at: str = None
    
    def __post_init__(self):
        if self.created_at is None:
            self.created_at = datetime.now().isoformat()


@dataclass
class Notebook:
    """笔记本"""
    id: str
    name: str
    description: str
    documents: List[SourceDocument]
    created_at: str
    updated_at: str
    tags: List[str]
    
    def to_dict(self) -> Dict:
        return {
            'id': self.id,
            'name': self.name,
            'description': self.description,
            'documents': [asdict(d) for d in self.documents],
            'created_at': self.created_at,
            'updated_at': self.updated_at,
            'tags': self.tags
        }


class DocumentProcessor:
    """文档处理器"""
    
    SUPPORTED_FORMATS = ['.pdf', '.txt', '.md', '.docx', '.html']
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # 下载NLTK数据
        try:
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            nltk.download('punkt', quiet=True)
        try:
            nltk.data.find('corpora/stopwords')
        except LookupError:
            nltk.download('stopwords', quiet=True)
    
    def process_file(self, filepath: str) -> SourceDocument:
        """处理文件"""
        path = Path(filepath)
        
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {filepath}")
        
        # 读取内容
        content = self._read_file(filepath)
        
        # 提取元数据
        metadata = self._extract_metadata(filepath, content)
        
        # 分块
        chunks = self._chunk_text(content)
        
        # 生成ID
        doc_id = hashlib.md5(f"{filepath}{datetime.now()}".encode()).hexdigest()[:12]
        
        return SourceDocument(
            id=doc_id,
            filename=path.name,
            filepath=str(path.absolute()),
            content=content,
            doc_type=path.suffix.lower(),
            metadata=metadata,
            chunks=chunks
        )
    
    def _read_file(self, filepath: str) -> str:
        """读取文件内容"""
        path = Path(filepath)
        suffix = path.suffix.lower()
        
        if suffix == '.pdf':
            return self._read_pdf(filepath)
        elif suffix == '.docx':
            return self._read_docx(filepath)
        elif suffix in ['.txt', '.md', '.html']:
            return self._read_text(filepath)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")
    
    def _read_pdf(self, filepath: str) -> str:
        """读取PDF"""
        text = []
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text.append(page.extract_text())
        return '\n'.join(text)
    
    def _read_docx(self, filepath: str) -> str:
        """读取Word文档"""
        doc = DocxDocument(filepath)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        return '\n'.join(text)
    
    def _read_text(self, filepath: str) -> str:
        """读取文本文件"""
        encodings = ['utf-8', 'gbk', 'latin-1']
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"无法解码文件: {filepath}")
    
    def _extract_metadata(self, filepath: str, content: str) -> Dict:
        """提取元数据"""
        path = Path(filepath)
        
        metadata = {
            'filename': path.name,
            'file_size': path.stat().st_size,
            'modified_time': datetime.fromtimestamp(path.stat().st_mtime).isoformat(),
            'word_count': len(content.split()),
            'char_count': len(content),
        }
        
        # 提取标题（假设第一行或前100字符）
        lines = content.strip().split('\n')
        if lines:
            metadata['title'] = lines[0][:100]
        
        # 提取关键词
        metadata['keywords'] = self._extract_keywords(content)
        
        return metadata
    
    def _extract_keywords(self, text: str, top_n: int = 10) -> List[str]:
        """提取关键词"""
        try:
            words = word_tokenize(text.lower())
            stop_words = set(stopwords.words('english'))
            
            # 过滤停用词和短词
            words = [w for w in words if w.isalnum() and len(w) > 2 and w not in stop_words]
            
            # 统计词频
            word_freq = Counter(words)
            return [word for word, _ in word_freq.most_common(top_n)]
        except:
            return []
    
    def _chunk_text(self, text: str) -> List[str]:
        """将文本分块"""
        chunks = []
        
        # 按句子分割
        sentences = sent_tokenize(text)
        
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            
            if current_length + sentence_length > self.chunk_size and current_chunk:
                # 保存当前块
                chunks.append(' '.join(current_chunk))
                
                # 保留重叠部分
                overlap_text = ' '.join(current_chunk[-2:]) if len(current_chunk) >= 2 else ''
                current_chunk = [overlap_text, sentence] if overlap_text else [sentence]
                current_length = len(overlap_text) + sentence_length
            else:
                current_chunk.append(sentence)
                current_length += sentence_length
        
        # 保存最后一块
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks


class NotebookLM:
    """NotebookLM主类"""
    
    def __init__(self, storage_dir: str = "~/.notebooklm"):
        self.storage_dir = Path(storage_dir).expanduser()
        self.storage_dir.mkdir(parents=True, exist_ok=True)
        
        self.processor = DocumentProcessor()
        self.notebooks: Dict[str, Notebook] = {}
        self.embedding_model = None
        
        # 加载嵌入模型
        if EMBEDDINGS_AVAILABLE:
            try:
                self.embedding_model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
                print("✅ 嵌入模型加载成功")
            except Exception as e:
                print(f"⚠️ 嵌入模型加载失败: {e}")
        
        # 加载已有笔记本
        self._load_notebooks()
    
    def _load_notebooks(self):
        """加载保存的笔记本"""
        notebooks_file = self.storage_dir / "notebooks.json"
        if notebooks_file.exists():
            try:
                with open(notebooks_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    # 这里简化处理，实际需要完整反序列化
                    print(f"📚 已加载 {len(data)} 个笔记本")
            except:
                pass
    
    def _save_notebooks(self):
        """保存笔记本"""
        notebooks_file = self.storage_dir / "notebooks.json"
        data = {nb_id: nb.to_dict() for nb_id, nb in self.notebooks.items()}
        with open(notebooks_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    
    def create_notebook(self, name: str, description: str = "") -> Notebook:
        """创建新笔记本"""
        notebook_id = hashlib.md5(f"{name}{datetime.now()}".encode()).hexdigest()[:12]
        
        notebook = Notebook(
            id=notebook_id,
            name=name,
            description=description,
            documents=[],
            created_at=datetime.now().isoformat(),
            updated_at=datetime.now().isoformat(),
            tags=[]
        )
        
        self.notebooks[notebook_id] = notebook
        self._save_notebooks()
        
        print(f"✅ 创建笔记本: {name} (ID: {notebook_id})")
        return notebook
    
    def add_document(self, notebook_id: str, filepath: str) -> SourceDocument:
        """添加文档到笔记本"""
        if notebook_id not in self.notebooks:
            raise ValueError(f"笔记本不存在: {notebook_id}")
        
        notebook = self.notebooks[notebook_id]
        
        # 处理文档
        print(f"📄 正在处理: {filepath}")
        document = self.processor.process_file(filepath)
        
        # 生成嵌入
        if self.embedding_model and document.chunks:
            print("🔢 生成文本嵌入...")
            document.embeddings = self.embedding_model.encode(document.chunks).tolist()
        
        notebook.documents.append(document)
        notebook.updated_at = datetime.now().isoformat()
        
        self._save_notebooks()
        
        print(f"✅ 已添加文档: {document.filename}")
        print(f"   字数: {document.metadata['word_count']}")
        print(f"   分块: {len(document.chunks)}")
        
        return document
    
    def generate_summary(self, notebook_id: str, style: str = "narrative") -> str:
        """生成摘要"""
        if notebook_id not in self.notebooks:
            raise ValueError(f"笔记本不存在: {notebook_id}")
        
        notebook = self.notebooks[notebook_id]
        
        if not notebook.documents:
            return "笔记本中没有文档"
        
        # 收集所有内容
        all_content = []
        for doc in notebook.documents:
            all_content.append(f"## {doc.filename}\n{doc.content[:2000]}...")
        
        combined_text = "\n\n".join(all_content)
        
        # 生成摘要（简化版，实际应调用AI模型）
        summary = self._generate_simple_summary(combined_text, style)
        
        return summary
    
    def _generate_simple_summary(self, text: str, style: str) -> str:
        """生成简单摘要（基于规则）"""
        sentences = sent_tokenize(text)
        
        if style == "bullet":
            # 提取关键句作为 bullet points
            key_sentences = sentences[:5]
            return "\n".join([f"• {s[:150]}..." for s in key_sentences])
        
        elif style == "executive":
            # 执行摘要
            word_count = len(text.split())
            return f"""# 执行摘要

本文档包含 {word_count} 个单词，主要涵盖以下内容：

{sentences[0][:200]}...

## 主要观点

1. {sentences[1][:150] if len(sentences) > 1 else 'N/A'}...
2. {sentences[2][:150] if len(sentences) > 2 else 'N/A'}...
3. {sentences[3][:150] if len(sentences) > 3 else 'N/A'}...

## 结论

{sentences[-1][:200] if sentences else 'N/A'}...
"""
        
        else:  # narrative
            return "\n".join(sentences[:10])
    
    def ask(self, notebook_id: str, question: str) -> str:
        """基于文档回答问题"""
        if notebook_id not in self.notebooks:
            raise ValueError(f"笔记本不存在: {notebook_id}")
        
        notebook = self.notebooks[notebook_id]
        
        if not notebook.documents:
            return "笔记本中没有文档"
        
        # 语义搜索相关片段
        relevant_chunks = self._semantic_search(notebook_id, question, top_k=3)
        
        if not relevant_chunks:
            return "未找到相关内容"
        
        # 构建回答（简化版）
        answer = f"基于文档内容，以下是相关信息：\n\n"
        for i, (chunk, score, doc_name) in enumerate(relevant_chunks, 1):
            answer += f"{i}. [{doc_name}] {chunk[:300]}...\n\n"
        
        return answer
    
    def _semantic_search(self, notebook_id: str, query: str, top_k: int = 3) -> List[Tuple[str, float, str]]:
        """语义搜索"""
        if not self.embedding_model:
            # 回退到关键词搜索
            return self._keyword_search(notebook_id, query, top_k)
        
        notebook = self.notebooks[notebook_id]
        
        # 编码查询
        query_embedding = self.embedding_model.encode([query])
        
        results = []
        for doc in notebook.documents:
            if doc.embeddings:
                # 计算相似度
                similarities = cosine_similarity(query_embedding, doc.embeddings)[0]
                
                # 获取最相似的块
                top_indices = similarities.argsort()[-top_k:][::-1]
                
                for idx in top_indices:
                    if idx < len(doc.chunks):
                        results.append((doc.chunks[idx], similarities[idx], doc.filename))
        
        # 排序并返回前k个
        results.sort(key=lambda x: x[1], reverse=True)
        return results[:top_k]
    
    def _keyword_search(self, notebook_id: str, query: str, top_k: int = 3) -> List[Tuple[str, float, str]]:
        """关键词搜索（回退方案）"""
        notebook = self.notebooks[notebook_id]
        query_words = set(query.lower().split())
        
        results = []
        for doc in notebook.documents:
            for chunk in doc.chunks:
                chunk_words = set(chunk.lower().split())
                overlap = len(query_words & chunk_words)
                score = overlap / len(query_words) if query_words else 0
                
                if score > 0:
                    results.append((chunk, score, doc.filename))
        
        results.sort(key=lambda x: x[1], reverse=True)
        return results[:top_k]
    
    def list_notebooks(self) -> List[Dict]:
        """列出所有笔记本"""
        return [
            {
                'id': nb.id,
                'name': nb.name,
                'description': nb.description,
                'document_count': len(nb.documents),
                'created_at': nb.created_at,
                'updated_at': nb.updated_at
            }
            for nb in self.notebooks.values()
        ]
    
    def get_notebook(self, notebook_id: str) -> Optional[Notebook]:
        """获取笔记本"""
        return self.notebooks.get(notebook_id)
    
    def delete_notebook(self, notebook_id: str):
        """删除笔记本"""
        if notebook_id in self.notebooks:
            del self.notebooks[notebook_id]
            self._save_notebooks()
            print(f"✅ 已删除笔记本: {notebook_id}")


# CLI接口
def main():
    """命令行入口"""
    import argparse
    
    parser = argparse.ArgumentParser(description='NotebookLM - 智能文档分析工具')
    subparsers = parser.add_subparsers(dest='command', help='可用命令')
    
    # 创建笔记本
    create_parser = subparsers.add_parser('create', help='创建笔记本')
    create_parser.add_argument('name', help='笔记本名称')
    create_parser.add_argument('--description', '-d', default='', help='描述')
    
    # 添加文档
    add_parser = subparsers.add_parser('add', help='添加文档')
    add_parser.add_argument('notebook_id', help='笔记本ID')
    add_parser.add_argument('files', nargs='+', help='文档文件路径')
    
    # 生成摘要
    summary_parser = subparsers.add_parser('summary', help='生成摘要')
    summary_parser.add_argument('notebook_id', help='笔记本ID')
    summary_parser.add_argument('--style', '-s', default='narrative', 
                               choices=['narrative', 'bullet', 'executive'],
                               help='摘要风格')
    
    # 提问
    ask_parser = subparsers.add_parser('ask', help='提问')
    ask_parser.add_argument('notebook_id', help='笔记本ID')
    ask_parser.add_argument('question', help='问题')
    
    # 列出现有笔记本
    list_parser = subparsers.add_parser('list', help='列出笔记本')
    
    args = parser.parse_args()
    
    if not args.command:
        parser.print_help()
        return
    
    # 初始化
    nblm = NotebookLM()
    
    if args.command == 'create':
        notebook = nblm.create_notebook(args.name, args.description)
        print(f"\n笔记本ID: {notebook.id}")
        print(f"使用此ID添加文档: notebooklm add {notebook.id} <文件>")
    
    elif args.command == 'add':
        for filepath in args.files:
            try:
                nblm.add_document(args.notebook_id, filepath)
            except Exception as e:
                print(f"❌ 添加失败: {e}")
    
    elif args.command == 'summary':
        try:
            summary = nblm.generate_summary(args.notebook_id, args.style)
            print("\n" + "="*60)
            print(summary)
            print("="*60)
        except Exception as e:
            print(f"❌ 生成摘要失败: {e}")
    
    elif args.command == 'ask':
        try:
            answer = nblm.ask(args.notebook_id, args.question)
            print("\n" + "="*60)
            print(answer)
            print("="*60)
        except Exception as e:
            print(f"❌ 回答问题失败: {e}")
    
    elif args.command == 'list':
        notebooks = nblm.list_notebooks()
        if notebooks:
            print("\n📚 笔记本列表:")
            print("-" * 80)
            for nb in notebooks:
                print(f"ID: {nb['id']}")
                print(f"名称: {nb['name']}")
                print(f"文档数: {nb['document_count']}")
                print(f"更新时间: {nb['updated_at']}")
                print("-" * 80)
        else:
            print("暂无笔记本")


if __name__ == '__main__':
    main()
