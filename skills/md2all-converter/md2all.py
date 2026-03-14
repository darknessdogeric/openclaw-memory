#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MD2ALL Converter - Markdown转PDF/Word/HTML全能转换器
支持中文，无需外部依赖
"""

import os
import sys
import re
from pathlib import Path

# 添加docx支持
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("正在安装 python-docx...")
    os.system("pip install python-docx -q")
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    import markdown
    HAS_MD = True
except ImportError:
    print("正在安装 markdown...")
    os.system("pip install markdown -q")
    import markdown

try:
    from fpdf import FPDF
    HAS_FPDF = True
except ImportError:
    print("正在安装 fpdf2...")
    os.system("pip install fpdf2 -q")
    from fpdf import FPDF

class MarkdownConverter:
    """Markdown转换器基类"""
    
    def __init__(self, md_file):
        self.md_file = Path(md_file)
        self.content = self._read_file()
        
    def _read_file(self):
        """读取Markdown文件"""
        with open(self.md_file, 'r', encoding='utf-8') as f:
            return f.read()
    
    def convert(self, output_file=None):
        """转换文件（子类实现）"""
        raise NotImplementedError

class MarkdownToWord(MarkdownConverter):
    """Markdown转Word转换器"""
    
    def __init__(self, md_file):
        super().__init__(md_file)
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """设置文档样式"""
        # 设置默认字体
        style = self.doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(11)
        
    def _parse_markdown(self):
        """解析Markdown内容"""
        lines = self.content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # 处理代码块
            if line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                self._add_code_block('\n'.join(code_lines))
                i += 1
                continue
            
            # 处理表格
            if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
                table_lines = [line]
                i += 1
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                self._add_table(table_lines)
                continue
            
            # 处理标题
            if line.startswith('# '):
                self._add_heading(line[2:], 1)
            elif line.startswith('## '):
                self._add_heading(line[3:], 2)
            elif line.startswith('### '):
                self._add_heading(line[4:], 3)
            elif line.startswith('#### '):
                self._add_heading(line[5:], 4)
            
            # 处理列表
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                self._add_bullet(line.strip()[2:])
            elif re.match(r'^\d+\.', line.strip()):
                self._add_numbered(line.strip())
            
            # 处理分隔线
            elif line.strip() == '---' or line.strip() == '***':
                self._add_horizontal_rule()
            
            # 处理普通段落
            elif line.strip():
                self._add_paragraph(line)
            
            # 空行
            else:
                self.doc.add_paragraph()
            
            i += 1
    
    def _add_heading(self, text, level):
        """添加标题"""
        heading = self.doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = 'Microsoft YaHei'
            run.font.bold = True
    
    def _add_paragraph(self, text):
        """添加段落"""
        # 处理加粗和斜体
        text = self._process_inline_formatting(text)
        p = self.doc.add_paragraph()
        self._add_formatted_text(p, text)
    
    def _process_inline_formatting(self, text):
        """处理行内格式标记"""
        # 暂时返回原文本，后续可实现完整Markdown解析
        return text
    
    def _add_formatted_text(self, paragraph, text):
        """添加带格式的文本"""
        # 处理加粗 **text**
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
        for part in parts:
            run = paragraph.add_run(part)
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(11)
            
            if part.startswith('**') and part.endswith('**'):
                run.text = part[2:-2]
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run.text = part[1:-1]
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run.text = part[1:-1]
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
    
    def _add_bullet(self, text):
        """添加无序列表"""
        p = self.doc.add_paragraph(style='List Bullet')
        self._add_formatted_text(p, text)
    
    def _add_numbered(self, text):
        """添加有序列表"""
        p = self.doc.add_paragraph(style='List Number')
        # 移除数字前缀
        text = re.sub(r'^\d+\.\s*', '', text)
        self._add_formatted_text(p, text)
    
    def _add_code_block(self, code):
        """添加代码块"""
        p = self.doc.add_paragraph()
        run = p.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.5)
    
    def _add_table(self, lines):
        """添加表格"""
        if len(lines) < 2:
            return
        
        # 解析表头
        headers = [cell.strip() for cell in lines[0].split('|') if cell.strip()]
        # 跳过分隔行
        # 解析数据行
        data_rows = []
        for line in lines[2:]:
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                data_rows.append(cells)
        
        if not headers:
            return
        
        # 创建表格
        table = self.doc.add_table(rows=1+len(data_rows), cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # 填充表头
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Microsoft YaHei'
        
        # 填充数据
        for i, row_data in enumerate(data_rows):
            row_cells = table.rows[i+1].cells
            for j, cell_text in enumerate(row_data):
                if j < len(row_cells):
                    row_cells[j].text = cell_text
                    for paragraph in row_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Microsoft YaHei'
    
    def _add_horizontal_rule(self):
        """添加水平分隔线"""
        p = self.doc.add_paragraph('_' * 50)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def convert(self, output_file=None):
        """转换为Word文档"""
        if output_file is None:
            output_file = self.md_file.with_suffix('.docx')
        
        self._parse_markdown()
        self.doc.save(output_file)
        print(f"[OK] Word文档已生成: {output_file}")
        return output_file

class MarkdownToPDF(MarkdownConverter):
    """Markdown转PDF转换器（使用FPDF）"""
    
    def __init__(self, md_file):
        super().__init__(md_file)
        self.pdf = FPDF()
        self.pdf.set_auto_page_break(auto=True, margin=15)
        self.pdf.add_page()
        self._setup_fonts()
    
    def _setup_fonts(self):
        """设置字体"""
        # 尝试加载中文字体
        try:
            # 尝试使用系统自带的中文字体
            font_paths = [
                "C:/Windows/Fonts/simhei.ttf",  # 黑体
                "C:/Windows/Fonts/simsun.ttc",  # 宋体
                "C:/Windows/Fonts/msyh.ttf",    # 微软雅黑
            ]
            
            for font_path in font_paths:
                if os.path.exists(font_path):
                    self.pdf.add_font('Chinese', '', font_path, uni=True)
                    self.pdf.add_font('Chinese', 'B', font_path, uni=True)
                    self.has_chinese_font = True
                    break
            else:
                self.has_chinese_font = False
                print("[WARN] 未找到中文字体，PDF中文可能显示为乱码")
        except Exception as e:
            print(f"⚠️ 字体加载失败: {e}")
            self.has_chinese_font = False
    
    def _parse_markdown(self):
        """解析并渲染Markdown"""
        lines = self.content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # 跳过代码块（简化处理）
            if line.startswith('```'):
                i += 1
                code_lines = []
                while i < len(lines) and not lines[i].startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                self._add_code_block('\n'.join(code_lines))
                i += 1
                continue
            
            # 处理标题
            if line.startswith('# '):
                self._add_heading(line[2:], 20)
            elif line.startswith('## '):
                self._add_heading(line[3:], 16)
            elif line.startswith('### '):
                self._add_heading(line[4:], 14)
            elif line.startswith('#### '):
                self._add_heading(line[5:], 12)
            
            # 处理列表
            elif line.strip().startswith('- ') or line.strip().startswith('* '):
                self._add_text('• ' + line.strip()[2:], 11)
            elif re.match(r'^\d+\.', line.strip()):
                self._add_text(line.strip(), 11)
            
            # 处理分隔线
            elif line.strip() == '---':
                self.pdf.ln(5)
                self.pdf.line(10, self.pdf.get_y(), 200, self.pdf.get_y())
                self.pdf.ln(5)
            
            # 处理普通段落
            elif line.strip():
                self._add_text(line, 11)
            
            # 空行
            else:
                self.pdf.ln(3)
            
            i += 1
    
    def _add_heading(self, text, size):
        """添加标题"""
        self.pdf.ln(5)
        if self.has_chinese_font:
            self.pdf.set_font('Chinese', 'B', size)
        else:
            self.pdf.set_font('Arial', 'B', size)
        self.pdf.set_x(10)
        self.pdf.multi_cell(190, 10, text)
        self.pdf.ln(2)
    
    def _add_text(self, text, size):
        """添加文本"""
        if not text.strip():
            return
            
        if self.has_chinese_font:
            self.pdf.set_font('Chinese', '', size)
        else:
            self.pdf.set_font('Arial', '', size)
        
        # 处理行宽，自动换行
        self.pdf.set_x(10)
        self.pdf.multi_cell(190, 6, text)
    
    def _add_code_block(self, code):
        """添加代码块"""
        self.pdf.ln(3)
        if self.has_chinese_font:
            self.pdf.set_font('Chinese', '', 9)
        else:
            self.pdf.set_font('Courier', '', 9)
        self.pdf.set_fill_color(240, 240, 240)
        self.pdf.set_x(15)
        self.pdf.multi_cell(180, 5, code, fill=True)
        self.pdf.ln(3)
    
    def convert(self, output_file=None):
        """转换为PDF"""
        if output_file is None:
            output_file = self.md_file.with_suffix('.pdf')
        
        self._parse_markdown()
        self.pdf.output(str(output_file))
        print(f"[OK] PDF文档已生成: {output_file}")
        return output_file

def convert_file(md_file, output_format='all'):
    """
    转换Markdown文件
    
    参数:
        md_file: Markdown文件路径
        output_format: 'pdf'|'docx'|'html'|'all'
    """
    md_path = Path(md_file)
    
    if not md_path.exists():
        print(f"[ERROR] 文件不存在: {md_file}")
        return
    
    results = []
    
    if output_format in ['docx', 'all']:
        try:
            converter = MarkdownToWord(md_file)
            docx_file = converter.convert()
            results.append(docx_file)
        except Exception as e:
            print(f"[ERROR] Word转换失败: {e}")
    
    if output_format in ['pdf', 'all']:
        try:
            converter = MarkdownToPDF(md_file)
            pdf_file = converter.convert()
            results.append(pdf_file)
        except Exception as e:
            print(f"[ERROR] PDF转换失败: {e}")
    
    if output_format in ['html', 'all']:
        try:
            with open(md_file, 'r', encoding='utf-8') as f:
                md_content = f.read()
            html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
            html_file = md_path.with_suffix('.html')
            with open(html_file, 'w', encoding='utf-8') as f:
                f.write(f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{md_path.stem}</title>
    <style>
        body {{ font-family: 'Microsoft YaHei', sans-serif; max-width: 900px; margin: 40px auto; padding: 20px; line-height: 1.6; }}
        h1 {{ color: #333; border-bottom: 2px solid #007acc; padding-bottom: 10px; }}
        h2 {{ color: #444; border-bottom: 1px solid #ddd; padding-bottom: 5px; margin-top: 30px; }}
        h3 {{ color: #555; margin-top: 25px; }}
        table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f5f5f5; font-weight: bold; }}
        tr:nth-child(even) {{ background-color: #fafafa; }}
        code {{ background-color: #f4f4f4; padding: 2px 5px; border-radius: 3px; font-family: Consolas, monospace; }}
        pre {{ background-color: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }}
        blockquote {{ border-left: 4px solid #007acc; margin: 0; padding-left: 20px; color: #666; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>""")
            print(f"[OK] HTML文档已生成: {html_file}")
            results.append(html_file)
        except Exception as e:
            print(f"[ERROR] HTML转换失败: {e}")
    
    return results

def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("""
MD2ALL Converter - Markdown全能转换器

用法:
    python md2all.py <markdown文件> [格式]
    
格式选项:
    pdf   - 转换为PDF
    docx  - 转换为Word文档
    html  - 转换为HTML
    all   - 转换全部格式（默认）

示例:
    python md2all.py README.md
    python md2all.py README.md pdf
    python md2all.py README.md docx
        """)
        return
    
    md_file = sys.argv[1]
    output_format = sys.argv[2] if len(sys.argv) > 2 else 'all'
    
    print(f"[*] 正在转换: {md_file}")
    print(f"[*] 输出格式: {output_format}")
    print("-" * 50)
    
    results = convert_file(md_file, output_format)
    
    if results:
        print("-" * 50)
        print(f"[DONE] 转换完成！共生成 {len(results)} 个文件")
        for f in results:
            print(f"   -> {f}")
    else:
        print("[FAIL] 转换失败")

if __name__ == '__main__':
    main()
