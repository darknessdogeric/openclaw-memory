#!/usr/bin/env python3
"""
文档内容提取器 - 支持 DOCX, PDF, TXT 等格式
"""

import sys
import os
import io

# 修复 Windows 控制台编码
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

def extract_docx(file_path):
    """提取 DOCX 文件内容"""
    try:
        from docx import Document
        doc = Document(file_path)
        
        content = []
        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    content.append(" | ".join(row_text))
        
        return "\n".join(content)
    except ImportError:
        return "ERROR: python-docx not installed. Run: pip install python-docx"
    except Exception as e:
        return f"ERROR extracting DOCX: {str(e)}"

def extract_pdf(file_path):
    """提取 PDF 文件内容"""
    try:
        import PyPDF2
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            content = []
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    content.append(f"--- Page {page_num + 1} ---\n{text}")
            return "\n\n".join(content)
    except ImportError:
        return "ERROR: PyPDF2 not installed. Run: pip install PyPDF2"
    except Exception as e:
        return f"ERROR extracting PDF: {str(e)}"

def extract_txt(file_path):
    """提取 TXT 文件内容"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # 尝试其他编码
        with open(file_path, 'r', encoding='gbk') as f:
            return f.read()
    except Exception as e:
        return f"ERROR reading TXT: {str(e)}"

def main():
    if len(sys.argv) < 2:
        print("Usage: python extract_doc.py <file_path>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    if not os.path.exists(file_path):
        print(f"ERROR: File not found: {file_path}")
        sys.exit(1)
    
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.docx':
        print(extract_docx(file_path))
    elif ext == '.pdf':
        print(extract_pdf(file_path))
    elif ext in ['.txt', '.md', '.json', '.xml', '.html', '.csv']:
        print(extract_txt(file_path))
    else:
        print(f"ERROR: Unsupported file format: {ext}")
        print("Supported: .docx, .pdf, .txt, .md, .json, .xml, .html, .csv")

if __name__ == "__main__":
    main()
