# -*- coding: utf-8 -*-
"""Markdown → Word 转换器（处理表格/标题/列表/粗体）"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

INPUT = r"C:\Users\Administrator\.openclaw\workspace\reports\2026-05-12-沙坪坝区五一酒店分析.md"
OUTPUT = r"C:\Users\Administrator\Desktop\2026五一沙坪坝区酒店分析报告.docx"

doc = Document()

# 页面设置
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# 样式预设
style_normal = doc.styles['Normal']
style_normal.font.name = '微软雅黑'
style_normal.font.size = Pt(10.5)
style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_heading_styled(text, level):
    """添加标题"""
    sizes = {1: Pt(22), 2: Pt(16), 3: Pt(13), 4: Pt(11)}
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if level <= 2:
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def add_paragraph_styled(text, bold=False, italic=False):
    """添加段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10.5)
    run.bold = bold
    run.italic = italic
    return p

def add_rich_paragraph(text):
    """解析内联格式（粗体、斜体）添加到段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    
    # 解析 **粗体** 和 *斜体*
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            # 处理行内代码
            codes = re.split(r'(`[^`]+`)', part)
            for c in codes:
                if c.startswith('`') and c.endswith('`'):
                    run = p.add_run(c[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                else:
                    run = p.add_run(c)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(10.5)
    return p

def add_table_from_md(lines, start_idx):
    """解析markdown表格并添加到doc"""
    # lines是表格行列表（含分隔线）
    rows_data = []
    for line in lines:
        if '|---' in line or '| ---' in line:
            continue  # 跳过分隔线
        cells = [c.strip() for c in line.split('|')[1:-1]]  # 去掉首尾空
        if cells:
            rows_data.append(cells)
    
    if len(rows_data) < 2:
        return
    
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row in enumerate(rows_data):
        for j, cell_text in enumerate(row):
            if j < len(table.rows[i].cells):
                cell = table.rows[i].cells[j]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.font.name = '微软雅黑'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(8.5)
                if i == 0:
                    run.bold = True
                    run.font.size = Pt(9)
    
    doc.add_paragraph()  # 表后空行

# 读取markdown
with open(INPUT, 'r', encoding='utf-8') as f:
    content = f.read()

lines = content.split('\n')

i = 0
in_table = False
table_lines = []
in_code_block = False
code_lines = []

while i < len(lines):
    line = lines[i]
    
    # 代码块处理
    if line.startswith('```'):
        if in_code_block:
            # 结束代码块
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(code_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            in_code_block = False
            code_lines = []
            i += 1
            continue
        else:
            in_code_block = True
            i += 1
            continue
    
    if in_code_block:
        code_lines.append(line)
        i += 1
        continue
    
    # 表格检测
    if line.startswith('|') and line.endswith('|'):
        table_lines.append(line)
        # 检测下一行是否还是表格
        if i + 1 < len(lines) and lines[i+1].startswith('|'):
            i += 1
            continue
        else:
            # 表格结束，渲染
            add_table_from_md(table_lines, i - len(table_lines) + 1)
            table_lines = []
            i += 1
            continue
    
    # 空行
    if line.strip() == '':
        i += 1
        continue
    
    # 标题
    if line.startswith('# '):
        add_heading_styled(line[2:], 1)
    elif line.startswith('## '):
        add_heading_styled(line[3:], 2)
    elif line.startswith('### '):
        add_heading_styled(line[4:], 3)
    elif line.startswith('#### '):
        add_heading_styled(line[5:], 4)
    
    # 水平线
    elif line.strip() == '---' or line.strip() == '***':
        doc.add_paragraph('─' * 50)
    
    # 引用
    elif line.startswith('> '):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line[2:])
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.size = Pt(10)
        run.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    # 无序列表
    elif line.startswith('- ') or line.startswith('* '):
        text = line[2:]
        add_rich_paragraph('• ' + text)
    
    # 有序列表
    elif re.match(r'^\d+\.\s', line):
        text = re.sub(r'^\d+\.\s', '', line)
        add_rich_paragraph(line.strip())
    
    # 普通段落
    else:
        # 处理内联格式
        text = line.strip()
        if text:
            add_rich_paragraph(text)
    
    i += 1

# 保存
doc.save(OUTPUT)
print(f"✅ 已保存: {OUTPUT}")
print(f"   文件大小: {__import__('os').path.getsize(OUTPUT):,} bytes")
