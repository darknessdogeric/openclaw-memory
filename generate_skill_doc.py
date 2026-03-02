# -*- coding: utf-8 -*-
"""
生成 AHL 场景SKILL化架构方案 Word文档
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='Microsoft YaHei', size=11, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(size)
    font.bold = bold
    if bold:
        font.color.rgb = RGBColor(0, 0, 0)

def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    set_chinese_font(run, size=(18 if level==1 else (14 if level==2 else 12)), bold=True)
    return heading

def add_paragraph_zh(doc, text, bold=False, size=11):
    """添加中文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_chinese_font(run, size=size, bold=bold)
    return p

def add_table_zh(doc, rows, cols, header_texts):
    """添加表格"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    
    # 设置表头
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(header_texts):
        if i < len(hdr_cells):
            run = hdr_cells[i].paragraphs[0].add_run(text)
            set_chinese_font(run, bold=True)
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return table

def create_skill_document():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # 标题
    title = doc.add_heading(level=0)
    run = title.add_run('AHL 场景SKILL化架构方案 V4.0')
    set_chinese_font(run, size=22, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    run = subtitle.add_run('全服务式高星级酒店细分SKILL体系')
    set_chinese_font(run, size=14, bold=True)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    meta = doc.add_paragraph()
    run = meta.add_run('架构理念: 大模型作为中枢大脑，双AGENT作为执行主体，下设可插拔式细分场景SKILL，实现"乐高式"组合与快速迭代')
    set_chinese_font(run, size=10)
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 一、整体架构升级
    add_heading_zh(doc, '一、整体架构升级', 1)
    
    add_paragraph_zh(doc, '【架构图说明】')
    add_paragraph_zh(doc, '┌─────────────────────────────────────────────────────────────┐')
    add_paragraph_zh(doc, '│                     HAL-Brain 住宿业垂类大模型               │')
    add_paragraph_zh(doc, '│                   (DeepSeek/Qwen + 全行业知识图谱)            │')
    add_paragraph_zh(doc, '└─────────────────────────────────────────────────────────────┘')
    add_paragraph_zh(doc, '                                │')
    add_paragraph_zh(doc, '        ┌───────────────────────┼───────────────────────┐')
    add_paragraph_zh(doc, '        ▼                       ▼                       ▼')
    add_paragraph_zh(doc, '┌───────────────┐      ┌───────────────┐      ┌───────────────┐')
    add_paragraph_zh(doc, '│  C端 AI管家    │      │  B端 AI运营官  │      │  SKILL市场    │')
    add_paragraph_zh(doc, '│(对客服务中枢)  │      │(对店管理中枢)  │      │(第三方开发)   │')
    add_paragraph_zh(doc, '└───────┬───────┘      └───────┬───────┘      └───────────────┘')
    add_paragraph_zh(doc, '        │                      │')
    add_paragraph_zh(doc, '   ┌────┴────┐            ┌────┴────┐')
    add_paragraph_zh(doc, '   ▼         ▼            ▼         ▼')
    add_paragraph_zh(doc, '[服务SKILL池]          [管理SKILL池]')
    add_paragraph_zh(doc, '(对客场景)              (对店场景)')
    
    doc.add_paragraph()
    add_paragraph_zh(doc, '【核心变化】', bold=True)
    add_paragraph_zh(doc, '1. SKILL模块化: 每个细分场景封装为独立SKILL，可插拔配置')
    add_paragraph_zh(doc, '2. 多层级调用: AGENT→SKILL→子SKILL→原子能力')
    add_paragraph_zh(doc, '3. 场景组合: 不同业态(商务/度假/康养)通过SKILL组合快速适配')
    add_paragraph_zh(doc, '4. 开放生态: 支持第三方开发SKILL，通过HAL认证后接入')
    
    doc.add_page_break()
    
    # 二、C端AI管家
    add_heading_zh(doc, '二、C端AI管家：对客服务SKILL矩阵', 1)
    
    # 2.1 客房服务SKILL群
    add_heading_zh(doc, '2.1 客房服务SKILL群 (Room Service Skills)', 2)
    
    table = add_table_zh(doc, 8, 4, ['SKILL名称', '功能描述', '触发场景', '输出结果'])
    data = [
        ['RS-01 智能预订', '需求解析→房型匹配→库存查询→支付闭环', '"我想订个安静的房间"', '推荐3个选项+一键支付'],
        ['RS-02 个性化入住', '偏好记忆→房间预配置→自助入住指引', '确认预订后', '定制化欢迎指南'],
        ['RS-03 智能房控', '语音/文字控制灯光/空调/窗帘', '"把空调调到24度"', '设备状态确认'],
        ['RS-04 客房服务', '物品需求→工单派发→进度追踪', '"我需要加一床被子"', '服务时长预估'],
        ['RS-05 快速退房', '账单核对→发票开具→押金退还', '"我要退房"', '电子账单+发票'],
        ['RS-06 延住处理', '房态检查→价格计算→续费确认', '"想多住一晚"', '续住方案+支付'],
        ['RS-07 换房升级', '需求评估→差价计算→房态协调', '"这个房间太吵"', '升级方案']
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i+1]
        for j, text in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(text)
            set_chinese_font(run)
    
    doc.add_paragraph()
    add_paragraph_zh(doc, '【SKILL交互示例: RS-01 智能预订】', bold=True)
    add_paragraph_zh(doc, '用户: "下周三出差，需要个离机场近的商务房，预算600左右"')
    add_paragraph_zh(doc, '↓')
    add_paragraph_zh(doc, '[RS-01 智能预订 SKILL]')
    add_paragraph_zh(doc, '├─ 意图解析: 商务差旅 | 时间: 下周三 | 位置: 机场附近 | 预算: ¥600')
    add_paragraph_zh(doc, '├─ 房型匹配: 调取"商务大床房"标签 → 筛选距离机场<30min')
    add_paragraph_zh(doc, '├─ 库存查询: 直连PMS API → 可用房型A/B/C')
    add_paragraph_zh(doc, '├─ 场景推荐: "根据您的航班时间，推荐行政楼层(含快速安检通道)"')
    add_paragraph_zh(doc, '└─ 支付闭环: 发送微信支付链接(含企业报销所需信息)')
    add_paragraph_zh(doc, '↓')
    add_paragraph_zh(doc, '输出: "为您找到3个选项：[方案A]行政大床¥598(含早+机场接送)..."')
    
    doc.add_page_break()
    
    # 2.2 餐饮服务SKILL群
    add_heading_zh(doc, '2.2 餐饮服务SKILL群 (F&B Service Skills)', 2)
    
    add_paragraph_zh(doc, '【中西餐零点SKILL】', bold=True)
    table = add_table_zh(doc, 5, 3, ['SKILL名称', '功能描述', '典型对话'])
    data = [
        ['FB-01 餐厅顾问', '场景需求→餐厅推荐→预订占位', '"适合约会的餐厅"'],
        ['FB-02 菜品推荐', '口味偏好→dietary restriction→菜品匹配', '"推荐清淡一点的菜"'],
        ['FB-03 营养搭配', '健康目标→热量计算→套餐建议', '"减脂期吃什么"'],
        ['FB-04 酒水搭配', '菜品类型→酒水知识→推荐酒单', '"牛排配什么酒"']
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i+1]
        for j, text in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(text)
            set_chinese_font(run)
    
    doc.add_paragraph()
    add_paragraph_zh(doc, '【中餐包间SKILL - 商务宴请场景】', bold=True)
    table = add_table_zh(doc, 5, 3, ['SKILL名称', '功能描述', '核心能力'])
    data = [
        ['FB-06 宴请策划', '人数/预算/规格→包间推荐→菜单设计', '商务礼仪知识'],
        ['FB-07 面子工程', '菜品档次→摆盘呈现→仪式感营造', '高端菜品知识库'],
        ['FB-08 酒水管家', '白酒/红酒/茶→品牌档次→文化讲解', '酒文化+品牌故事'],
        ['FB-09 场景布置', '商务/庆生/节日→氛围方案→增值服务', '花艺+灯光+音乐']
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i+1]
        for j, text in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(text)
            set_chinese_font(run)
    
    doc.add_paragraph()
    add_paragraph_zh(doc, '【SKILL交互示例: FB-06 宴请策划】', bold=True)
    add_paragraph_zh(doc, '用户: "周五晚上商务宴请，6个人，人均500，要个有面子的包间"')
    add_paragraph_zh(doc, '↓')
    add_paragraph_zh(doc, '[FB-06 宴请策划 SKILL]')
    add_paragraph_zh(doc, '├─ 需求解析: 商务性质 | 6人 | ¥500/人 | 高端定位')
    add_paragraph_zh(doc, '├─ 包间匹配: 调取"翡翠厅"(12人圆桌+独立茶歇区+智能灯光)')
    add_paragraph_zh(doc, '├─ 菜单设计: 古法蒸东星斑+黑椒澳洲和牛+手工虾饺')
    add_paragraph_zh(doc, '├─ 面子加成: "本月新装修完，全市唯一智能灯光包间"')
    add_paragraph_zh(doc, '└─ 增值服务: 赠送果盘+每人伴手礼')
    
    doc.add_page_break()
    
    # 2.3 宴会服务SKILL群
    add_heading_zh(doc, '2.3 宴会服务SKILL群 (Events & Banquet Skills)', 2)
    
    table = add_table_zh(doc, 7, 4, ['SKILL名称', '适用场景', '功能', 'LTV运营'])
    data = [
        ['EV-01 婚宴顾问', '婚礼筹备', '全流程策划+供应商协调', '回门宴+宝宝宴'],
        ['EV-02 寿宴策划', '长辈庆生', '传统礼仪+家庭关系维护', '家庭聚餐'],
        ['EV-03 宝宝宴服务', '满月/百日/周岁', '成长节点运营+复购预埋', '成长链'],
        ['EV-04 商务宴会', '企业年会/答谢', 'B2B关系维护+方案定制', '长期协议'],
        ['EV-05 乔迁/升学宴', '家庭喜庆', '场景设计+亲友互动', '社群裂变'],
        ['EV-06 战友/同学会', '怀旧聚会', '情感连接+纪念服务', 'N次转介绍']
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i+1]
        for j, text in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(text)
            set_chinese_font(run)
    
    doc.add_paragraph()
    add_paragraph_zh(doc, '【EV-01 婚宴顾问SKILL - 子模块结构】', bold=True)
    add_paragraph_zh(doc, '[EV-01 婚宴顾问]')
    add_paragraph_zh(doc, '├─ [子SKILL] 需求诊断: 预算区间 | 桌数 | 风格偏好 | 特殊需求')
    add_paragraph_zh(doc, '├─ [子SKILL] 厅房匹配: 容量计算 | 动线设计 | 布置潜力评估')
    add_paragraph_zh(doc, '├─ [子SKILL] 菜单设计: 档次匹配 | 寓意菜品 | dietary restriction')
    add_paragraph_zh(doc, '├─ [子SKILL] 流程策划: 仪式流程 | 时间节点 | 应急预案')
    add_paragraph_zh(doc, '├─ [子SKILL] 供应商整合: 婚庆 | 摄影 | 化妆 | 司仪 推荐与协调')
    add_paragraph_zh(doc, '└─ [子SKILL] 复购预埋: 回门宴推荐 | 宝宝宴铺垫 | 周年纪念提醒')
    
    doc.add_page_break()
    
    # 2.4 前厅礼宾SKILL群
    add_heading_zh(doc, '2.4 前厅礼宾SKILL群 (Concierge Skills)', 2)
    
    table = add_table_zh(doc, 7, 3, ['SKILL名称', '功能描述', '知识库'])
    data = [
        ['CO-01 本地向导', '景点/餐厅/购物推荐', '周边POI+实时人流'],
        ['CO-02 交通管家', '接送机/租车/代驾安排', '车队资源+导航API'],
        ['CO-03 票务服务', '机票/火车票/景点票代购', '票务接口+退改政策'],
        ['CO-04 商务中心', '打印/翻译/会议室预订', '设备状态+翻译资源'],
        ['CO-05 应急支援', '医疗/报警/领事协助', '应急联系人+流程SOP'],
        ['CO-06 特殊关怀', 'VIP识别/纪念日/惊喜服务', 'CRM数据+创意库']
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i+1]
        for j, text in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(text)
            set_chinese_font(run)
    
    # 2.5 第四空间零售SKILL群
    add_heading_zh(doc, '2.5 第四空间零售SKILL群 (Fourth Space Retail Skills)', 2)
    
    table = add_table_zh(doc, 6, 4, ['SKILL名称', '场景触发', '转化路径', '客单价'])
    data = [
        ['RT-01 睡眠场景导购', '"枕头很舒服"', '扫码→详情页→下单→邮寄', '¥199-2999'],
        ['RT-02 洗浴场景导购', '"沐浴露什么牌子"', '成分介绍→购买链接→复购提醒', '¥49-199'],
        ['RT-03 茶饮场景导购', '"这茶很好喝"', '产地故事→订阅服务→月度配送', '¥39-299'],
        ['RT-04 阅读场景导购', '"这本书哪里买"', '书单推荐→购买链接→读书会邀请', '¥39-128'],
        ['RT-05 穿搭场景导购', '"睡衣很舒服"', '尺码推荐→套装搭配→会员折扣', '¥199-499']
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i+1]
        for j, text in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(text)
            set_chinese_font(run)
    
    doc.add_page_break()
    
    # 三、B端AI运营官
    add_heading_zh(doc, '三、B端AI运营官：对店管理SKILL矩阵', 1)
    
    # 3.1 收益管理SKILL群
    add_heading_zh(doc, '3.1 收益管理SKILL群 (Revenue Management Skills)', 2)
    
    table = add_table_zh(doc, 6, 4, ['SKILL名称', '功能', '触发条件', '执行动作'])
    data = [
        ['RM-01 动态定价', '竞品监测→需求预测→价格调整', '每6小时/市场异动', 'PMS价格自动更新'],
        ['RM-02 库存优化', '入住率预测→超售策略→房型升级', '入住前7天/满房', '自动释放/关闭渠道'],
        ['RM-03 渠道管理', 'OTA比价→佣金优化→直客引导', '每日', '价差告警+调价建议'],
        ['RM-04 细分市场定价', 'B2B协议价→会员价→walk-in价', '协议到期/新客', '分层定价策略'],
        ['RM-05 餐饮收益', '餐位利用率→动态菜单→套餐优化', '每日', '推荐菜调整+促销']
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i+1]
        for j, text in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(text)
            set_chinese_font(run)
    
    doc.add_paragraph()
    add_paragraph_zh(doc, '【RM-01 动态定价SKILL - 决策逻辑】', bold=True)
    add_paragraph_zh(doc, '每6小时触发 → ')
    add_paragraph_zh(doc, '├─ 数据采集: 本店数据+竞品价格+市场信号+渠道转化')
    add_paragraph_zh(doc, '├─ 需求预测 (ML模型): 输入多维度特征 → 输出未来14天需求概率')
    add_paragraph_zh(doc, '├─ 定价决策: 高需求日(>80%)上调10-20% | 低需求日(<50%)促销')
    add_paragraph_zh(doc, '├─ 价格发布: PMS更新 + OTA同步 + 私域差异化定价')
    add_paragraph_zh(doc, '└─ 效果追踪: 转化率监控 → 强化学习反馈')
    
    doc.add_page_break()
    
    # 其他B端SKILL群
    add_heading_zh(doc, '3.2 前厅运营SKILL群', 2)
    table = add_table_zh(doc, 6, 3, ['SKILL名称', '功能', '适用场景'])
    data = [
        ['FO-01 智能排房', '客史偏好→房态→最优分配', '每日入住前'],
        ['FO-02 夜班值守', '22:00-08:00自动应答+工单', '夜间'],
        ['FO-03 客诉处理', '情绪识别→分级响应→解决方案', '投诉触发'],
        ['FO-04 VIP识别', 'CRM标签→个性化服务预案', '预订/入住'],
        ['FO-05 失物管理', '登记→寻回→邮寄/保管', '离店后']
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i+1]
        for j, text in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(text)
            set_chinese_font(run)
    
    add_heading_zh(doc, '3.3 客房管理SKILL群', 2)
    table = add_table_zh(doc, 6, 3, ['SKILL名称', '功能', '触发条件'])
    data = [
        ['HK-01 智能排班', '入住率→清洁量→人力排班', '每日早班前'],
        ['HK-02 客房检查', '退房→查房→损耗标记→计费', '每次退房'],
        ['HK-03 布草管理', '库存预警→换洗周期→成本分析', '实时/周度'],
        ['HK-04 维保触发', '设备异常→工单派发→跟踪', 'IoT告警/客报'],
        ['HK-05 Minibar补货', '消耗统计→补货清单→库存优化', '每日']
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i+1]
        for j, text in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(text)
            set_chinese_font(run)
    
    add_heading_zh(doc, '3.4 餐饮运营SKILL群', 2)
    table = add_table_zh(doc, 5, 3, ['SKILL名称', '功能', '输出'])
    data = [
        ['FB-OP-01 餐位管理', '预订→排位→翻台优化', '实时等位信息'],
        ['FB-OP-02 菜单工程', '成本率→点击率→毛利分析', '推荐菜调整建议'],
        ['FB-OP-03 库存预警', '食材消耗→保质期→补货提醒', '采购清单'],
        ['FB-OP-04 厨师排班', '预订量→人力需求→排班表', '优化排班方案']
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i+1]
        for j, text in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(text)
            set_chinese_font(run)
    
    doc.add_page_break()
    
    add_heading_zh(doc, '3.5 宴会销售SKILL群', 2)
    table = add_table_zh(doc, 6, 4, ['SKILL名称', '功能', '转化漏斗', 'LTV运营'])
    data = [
        ['BS-01 婚宴销售', '线索→跟进→签约→执行→复购', '回门宴+宝宝宴', '成长链'],
        ['BS-02 寿宴销售', '家庭关系挖掘→年度聚会', '家庭聚餐', '家庭LTV'],
        ['BS-03 宝宝宴销售', '成长节点追踪→持续运营', '成长链', '年度复购'],
        ['BS-04 商务宴销售', '企业关系维护→年会/日常', 'B2B长期协议', '企业LTV'],
        ['BS-05 同学/战友会', '社群渗透→裂变转介绍', 'N次转介绍', '裂变获客']
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i+1]
        for j, text in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(text)
            set_chinese_font(run)
    
    add_heading_zh(doc, '3.6 闲置空间运营SKILL群', 2)
    table = add_table_zh(doc, 7, 3, ['SKILL名称', '空间类型', '运营模式'])
    data = [
        ['SP-01 会议室运营', '会议室', '时段租赁+企业服务'],
        ['SP-02 包间二次利用', '餐厅包间', '非餐期商务洽谈/棋牌'],
        ['SP-03 大堂吧激活', '大堂吧', '快闪店/咖啡合作'],
        ['SP-04 户外空间', '花园/露台', '市集/婚礼/派对'],
        ['SP-05 停车场增值', '停车场', '夜间市集/充电桩'],
        ['SP-06 屋顶开发', '屋顶/露台', '酒吧/观星/瑜伽']
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i+1]
        for j, text in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(text)
            set_chinese_font(run)
    
    doc.add_page_break()
    
    add_heading_zh(doc, '3.7 B2B企业市场SKILL群', 2)
    table = add_table_zh(doc, 6, 3, ['SKILL名称', '功能', '决策链'])
    data = [
        ['B2B-01 RFP响应', '招标信息→智能匹配→自动报价', '企业采购'],
        ['B2B-02 TMC对接', '协议价同步→库存管理→结算对账', '差旅管理公司'],
        ['B2B-03 MICE销售', '场地匹配→方案打包→执行协调', '会议策划者'],
        ['B2B-04 协议客维护', '企业客史→消费分析→续约提醒', '企业行政'],
        ['B2B-05 长包房管理', '月租客户→服务定制→续租预测', '长住客']
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i+1]
        for j, text in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(text)
            set_chinese_font(run)
    
    add_heading_zh(doc, '3.8 财务自动化SKILL群', 2)
    table = add_table_zh(doc, 6, 3, ['SKILL名称', '功能', '效率提升'])
    data = [
        ['FI-01 自动对账', 'PMS+OTA+支付三方对账', '节省80%人力'],
        ['FI-02 收入审计', '房费+餐饮+其他收入核验', '差错自动标红'],
        ['FI-03 成本分析', '部门成本→预算对比→预警', '实时成本看板'],
        ['FI-04 报表生成', '日报/月报/年报自动生成', '一键导出'],
        ['FI-05 税务合规', '发票管理→税务申报→风险预警', '合规自动化']
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i+1]
        for j, text in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(text)
            set_chinese_font(run)
    
    add_heading_zh(doc, '3.9 能耗管理SKILL群', 2)
    table = add_table_zh(doc, 5, 3, ['SKILL名称', '功能', '节能效果'])
    data = [
        ['EN-01 智能温控', '入住率→空调策略→温度调节', '节能15-25%'],
        ['EN-02 照明优化', '自然光+人感→灯光控制', '节能10-15%'],
        ['EN-03 用水监控', '流量监测→异常告警→节水建议', '节水10-20%'],
        ['EN-04 能耗报表', '部门能耗→同比环比→优化建议', '可视化分析']
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i+1]
        for j, text in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(text)
            set_chinese_font(run)
    
    doc.add_page_break()
    
    # 四、业态适配
    add_heading_zh(doc, '四、SKILL编排与业态适配', 1)
    
    add_heading_zh(doc, '4.1 全服务高星级酒店标准配置', 2)
    add_paragraph_zh(doc, '【全服务高星级酒店 - 标准配置包】', bold=True)
    add_paragraph_zh(doc, '├── C端AI管家 (必备SKILL)')
    add_paragraph_zh(doc, '│   ├── 客房服务SKILL群 [7个]')
    add_paragraph_zh(doc, '│   ├── 餐饮服务SKILL群 [13个]')
    add_paragraph_zh(doc, '│   ├── 宴会服务SKILL群 [6个]')
    add_paragraph_zh(doc, '│   ├── 前厅礼宾SKILL群 [6个]')
    add_paragraph_zh(doc, '│   └── 零售导购SKILL群 [5个]')
    add_paragraph_zh(doc, '│')
    add_paragraph_zh(doc, '├── B端AI运营官 (必备SKILL)')
    add_paragraph_zh(doc, '│   ├── 收益管理SKILL群 [5个]')
    add_paragraph_zh(doc, '│   ├── 前厅运营SKILL群 [5个]')
    add_paragraph_zh(doc, '│   ├── 客房管理SKILL群 [5个]')
    add_paragraph_zh(doc, '│   ├── 餐饮运营SKILL群 [8个]')
    add_paragraph_zh(doc, '│   ├── 宴会销售SKILL群 [5个]')
    add_paragraph_zh(doc, '│   ├── 市场营销SKILL群 [5个]')
    add_paragraph_zh(doc, '│   ├── 空间运营SKILL群 [6个]')
    add_paragraph_zh(doc, '│   ├── B2B市场SKILL群 [5个]')
    add_paragraph_zh(doc, '│   ├── 财务自动化SKILL群 [5个]')
    add_paragraph_zh(doc, '│   └── 能耗管理SKILL群 [4个]')
    add_paragraph_zh(doc, '│')
    add_paragraph_zh(doc, '└── 扩展包 (按需选配)')
    add_paragraph_zh(doc, '    ├── 康养长住SKILL群 [5个] (康养酒店)')
    add_paragraph_zh(doc, '    ├── 公寓运营SKILL群 [4个] (服务式公寓)')
    add_paragraph_zh(doc, '    └── MICE深度SKILL群 [3个] (会议型酒店)')
    
    add_heading_zh(doc, '4.2 不同业态SKILL组合建议', 2)
    table = add_table_zh(doc, 7, 4, ['业态类型', '核心SKILL', '差异化SKILL', '禁用SKILL'])
    data = [
        ['商务酒店', '客房+收益+B2B', '快速入住+会议服务', '宴会(小型)'],
        ['度假酒店', '客房+礼宾+零售', '本地向导+活动预订', '长包房管理'],
        ['康养酒店', '客房+康养+长住', '健康档案+就医协助', '高频换房'],
        ['服务公寓', '长住+B2B+自助', '月租管理+企业协议', '宴会(大型)'],
        ['精品民宿', '客房+礼宾+零售', 'IP打造+内容生成', 'B2B(标准)'],
        ['会议酒店', '宴会+MICE+空间', 'RFP响应+会场管理', '零售(弱化)']
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i+1]
        for j, text in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(text)
            set_chinese_font(run)
    
    doc.add_page_break()
    
    # 五、技术实现
    add_heading_zh(doc, '五、技术实现要点', 1)
    
    add_heading_zh(doc, '5.1 SKILL封装标准', 2)
    add_paragraph_zh(doc, '每个SKILL需包含以下标准结构:')
    add_paragraph_zh(doc, '')
    add_paragraph_zh(doc, 'metadata:')
    add_paragraph_zh(doc, '  name: "智能预订"')
    add_paragraph_zh(doc, '  code: "RS-01"')
    add_paragraph_zh(doc, '  version: "1.0.0"')
    add_paragraph_zh(doc, '  category: "客房服务"')
    add_paragraph_zh(doc, '')
    add_paragraph_zh(doc, 'trigger:')
    add_paragraph_zh(doc, '  keywords: ["订房", "预订", "房间"]')
    add_paragraph_zh(doc, '  intents: ["预订意图", "查询意图"]')
    add_paragraph_zh(doc, '')
    add_paragraph_zh(doc, 'input:')
    add_paragraph_zh(doc, '  required: ["入住日期", "离店日期"]')
    add_paragraph_zh(doc, '  optional: ["预算", "偏好", "特殊需求"]')
    add_paragraph_zh(doc, '')
    add_paragraph_zh(doc, 'process:')
    add_paragraph_zh(doc, '  steps: [意图解析, 需求提取, 库存查询, 推荐生成, 支付闭环]')
    add_paragraph_zh(doc, '')
    add_paragraph_zh(doc, 'output:')
    add_paragraph_zh(doc, '  type: "卡片+按钮"')
    add_paragraph_zh(doc, '  fields: ["推荐房型", "价格", "图片", "支付链接"]')
    add_paragraph_zh(doc, '')
    add_paragraph_zh(doc, 'fallback:')
    add_paragraph_zh(doc, '  condition: "无房/系统故障"')
    add_paragraph_zh(doc, '  action: "转人工+补偿方案"')
    add_paragraph_zh(doc, '')
    add_paragraph_zh(doc, 'metrics:')
    add_paragraph_zh(doc, '  success_rate: ">90%"')
    add_paragraph_zh(doc, '  avg_response_time: "<2s"')
    
    add_heading_zh(doc, '5.2 SKILL调用链示例', 2)
    add_paragraph_zh(doc, '用户输入: "我想订个安静能看到海的房间，下周三入住"')
    add_paragraph_zh(doc, '   ↓')
    add_paragraph_zh(doc, '[AGENT路由]')
    add_paragraph_zh(doc, '├─ 意图识别: 预订客房')
    add_paragraph_zh(doc, '├─ 情感分析: 中性+期待')
    add_paragraph_zh(doc, '└─ SKILL选择: RS-01 智能预订')
    add_paragraph_zh(doc, '   ↓')
    add_paragraph_zh(doc, '[RS-01 智能预订 SKILL]')
    add_paragraph_zh(doc, '├─ [子步骤] 需求解析')
    add_paragraph_zh(doc, '│   ├─ 时间: 下周三(2026-03-11) | 默认1晚')
    add_paragraph_zh(doc, '│   ├─ 偏好: 安静(噪音<40dB) | 海景(海景房标签)')
    add_paragraph_zh(doc, "│   └─ 预算: 未指定(全价位)")
    add_paragraph_zh(doc, '│')
    add_paragraph_zh(doc, '├─ [子步骤] 库存查询 (调用PMS API)')
    add_paragraph_zh(doc, '│   ├─ 房型A: 海景大床房 | ¥698 | 剩余3间 | 东向(安静)')
    add_paragraph_zh(doc, '│   ├─ 房型B: 行政海景房 | ¥998 | 剩余1间 | 高层+行政待遇')
    add_paragraph_zh(doc, '│   └─ 房型C: 豪华海景套 | ¥1688 | 剩余1间 | 露台+客厅')
    add_paragraph_zh(doc, '│')
    add_paragraph_zh(doc, '├─ [子步骤] 场景推荐 (调用知识图谱)')
    add_paragraph_zh(doc, '│   └─ "您选的日期正值樱花季，推荐房型B(行政楼层)，')
    add_paragraph_zh(doc, '│       含行政酒廊下午茶，可在18楼俯瞰海景+樱花"')
    add_paragraph_zh(doc, '│')
    add_paragraph_zh(doc, '└─ [子步骤] 支付闭环 (调用支付SKILL)')
    add_paragraph_zh(doc, '    └─ 生成支付链接 → 发送给用户')
    add_paragraph_zh(doc, '       ↓')
    add_paragraph_zh(doc, '输出: 卡片展示3个选项+一键支付')
    
    # 六、总结
    doc.add_page_break()
    add_heading_zh(doc, '六、总结与下一步', 1)
    
    add_heading_zh(doc, '6.1 核心价值', 2)
    table = add_table_zh(doc, 6, 3, ['维度', '传统方案', 'SKILL化方案'])
    data = [
        ['灵活性', '功能固定', '按需组合'],
        ['迭代速度', '月级更新', '周级更新'],
        ['个性化', '通用方案', '业态定制'],
        ['扩展性', '封闭系统', '开放生态'],
        ['学习进化', '规则驱动', '数据驱动']
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i+1]
        for j, text in enumerate(row_data):
            run = row.cells[j].paragraphs[0].add_run(text)
            set_chinese_font(run)
    
    add_heading_zh(doc, '6.2 实施建议', 2)
    add_paragraph_zh(doc, 'Phase 1 (1-2月): 核心SKILL开发', bold=True)
    add_paragraph_zh(doc, '• 客房预订+基础服务SKILL')
    add_paragraph_zh(doc, '• 收益管理+前厅运营SKILL')
    add_paragraph_zh(doc, '• 单店跑通闭环')
    add_paragraph_zh(doc, '')
    add_paragraph_zh(doc, 'Phase 2 (3-6月): 场景SKILL扩展', bold=True)
    add_paragraph_zh(doc, '• 餐饮+宴会SKILL群')
    add_paragraph_zh(doc, '• B2B+空间运营SKILL')
    add_paragraph_zh(doc, '• 业态差异化配置')
    add_paragraph_zh(doc, '')
    add_paragraph_zh(doc, 'Phase 3 (7-12月): 生态开放', bold=True)
    add_paragraph_zh(doc, '• SKILL市场上线')
    add_paragraph_zh(doc, '• 第三方开发者接入')
    add_paragraph_zh(doc, '• 行业标准化推进')
    
    doc.add_paragraph()
    add_paragraph_zh(doc, '---', bold=True)
    add_paragraph_zh(doc, '下一步行动:', bold=True)
    add_paragraph_zh(doc, '1. 确认SKILL优先级(建议从客房预订+收益管理开始)')
    add_paragraph_zh(doc, '2. 选取1家样板店进行SKILL化改造')
    add_paragraph_zh(doc, '3. 制定SKILL开发排期与测试计划')
    
    # 保存文档
    output_path = r"C:\Users\Administrator\Desktop\张实项目总控\06-AHL-去中心化旅行平台\AHL场景SKILL化架构方案V4.0.docx"
    doc.save(output_path)
    print(f"文档已生成: {output_path}")
    return output_path

if __name__ == "__main__":
    create_skill_document()