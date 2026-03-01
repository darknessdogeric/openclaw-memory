#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to DOCX Converter
Converts AHL product scheme MD to Word document with formatting
"""

import re
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def convert_md_to_docx(input_path, output_path):
    print(f"Reading: {input_path}")
    
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set default font for document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10.5)
    
    # Process lines
    lines = content.split('\n')
    i = 0
    in_table = False
    table_lines = []
    in_code = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Code blocks
        if line.strip().startswith('```'):
            if in_code:
                # End code block
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(50, 50, 50)
                in_code = False
                code_lines = []
            else:
                in_code = True
            i += 1
            continue
        
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        
        # Tables
        if '|' in line:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1            
            # Check if next line continues table
            if i < len(lines) and '|' not in lines[i]:
                # End of table
                in_table = False
                if len(table_lines) >= 2:  # Need at least header + separator
                    render_table(doc, table_lines)
                table_lines = []
            continue
        elif in_table:
            in_table = False
            if len(table_lines) >= 2:
                render_table(doc, table_lines)
            table_lines = []
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(26, 82, 118)
                run.font.size = Pt(24)
                run.font.bold = True
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=1)
            for run in p.runs:
                run.font.color.rgb = RGBColor(40, 116, 166)
                run.font.size = Pt(18)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=2)
            for run in p.runs:
                run.font.color.rgb = RGBColor(46, 134, 193)
                run.font.size = Pt(14)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=3)
            for run in p.runs:
                run.font.color.rgb = RGBColor(84, 153, 199)
                run.font.size = Pt(12)
        
        # Lists
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            text = clean_markdown(text)
            p = doc.add_paragraph(text, style='List Bullet')
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            text = clean_markdown(text)
            p = doc.add_paragraph(text, style='List Number')
        
        # Blockquote
        elif line.strip().startswith('> '):
            text = line.strip()[2:]
            text = clean_markdown(text)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(80, 80, 80)
            p.paragraph_format.left_indent = Inches(0.3)
        
        # Horizontal rule
        elif line.strip() in ['---', '***', '___']:
            doc.add_paragraph('_' * 50)
        
        # Empty line
        elif not line.strip():
            pass  # Skip empty lines
        
        # Regular paragraph
        else:
            text = clean_markdown(line)
            if text:
                p = doc.add_paragraph(text)
        
        i += 1
    
    # Handle remaining table
    if in_table and len(table_lines) >= 2:
        render_table(doc, table_lines)
    
    # Save document
    print(f"Saving: {output_path}")
    doc.save(output_path)
    print("Done!")
    return output_path

def clean_markdown(text):
    """Remove markdown formatting for plain text"""
    # Bold
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Italic
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # Code
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text.strip()

def render_table(doc, table_lines):
    """Render markdown table to Word table"""
    if len(table_lines) < 2:
        return
    
    # Parse header
    headers = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
    
    # Skip separator line
    rows = []
    for line in table_lines[2:]:
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells:
            rows.append(cells)
    
    if not headers or not rows:
        return
    
    # Create table
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = clean_markdown(header)
        set_cell_shading(hdr_cells[i], '3498db')
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for i, cell_text in enumerate(row_data):
            if i < len(row_cells):
                row_cells[i].text = clean_markdown(cell_text)
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
    
    doc.add_paragraph()  # Space after table

if __name__ == '__main__':
    input_file = r"C:\Users\Administrator\Desktop\张实项目总控\06-AHL-去中心化旅行平台\AHL通用产品方案_V3_专业版.md"
    output_file = r"C:\Users\Administrator\Desktop\张实项目总控\06-AHL-去中心化旅行平台\AHL通用产品方案_V3_专业版.docx"
    
    convert_md_to_docx(input_file, output_file)
