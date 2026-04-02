# -*- coding: utf-8 -*-
"""
AHL项目OPC路演PPT大纲 - Word文档生成
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

def set_run_font(run, font_name_cn='微软雅黑', font_name_en='Arial', size=11, bold=False):
    run.font.name = font_name_en
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)

def add_heading(doc, text, level=1, bold=True, size=16, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = '微软雅黑'
        run.font.size = Pt(size)
        run.font.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_paragraph(doc, text, bold=False, size=11, indent=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_bullet(doc, text, level=0, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(level * 0.75 + 0.5)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_table_row(table, col1, col2='', bold1=False, bold2=False):
    row = table.add_row()
    c1 = row.cells[0]
    c2 = row.cells[1]
    r1 = c1.paragraphs[0].add_run(col1)
    r1.font.bold = bold1
    r1.font.name = '微软雅黑'
    r1.font.size = Pt(10)
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if col2:
        r2 = c2.paragraphs[0].add_run(col2)
        r2.font.bold = bold2
        r2.font.name = '微软雅黑'
        r2.font.size = Pt(10)
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 创建文档
doc = Document()

# 设置页面边距
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# ==================== 封面 ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AHL项目OPC路演PPT大纲')
run.font.name = '微软雅黑'
run.font.size = Pt(24)
run.font.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('（6分钟精华版）')
run.font.name = '微软雅黑'
run.font.size = Pt(14)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()

# ==================== 整体结构 ====================
add_heading(doc, '一、整体结构', level=1, size=16)

# 表格
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

# 表头
hdr_cells = table.rows[0].cells
for i, text in enumerate(['页码', '标题', '时间']):
    hdr_cells[i].paragraphs[0].clear()
    run = hdr_cells[i].paragraphs[0].add_run(text)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 表格数据
data = [
    ('1', '封面', '10秒'),
    ('2', '我们在做什么', '40秒'),
    ('3', '我们要实现什么', '40秒'),
    ('4', '我们的愿景', '30秒'),
    ('5', '行业痛点', '50秒'),
    ('6', '解决方案', '50秒'),
    ('7', '产品架构', '50秒'),
    ('8', '技术壁垒', '40秒'),
    ('9', '商业模式', '40秒'),
    ('10', '市场机会', '30秒'),
    ('11', '团队优势', '30秒'),
    ('12', '融资需求', '30秒'),
]

for row_data in data:
    row = table.add_row()
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        run = cell.paragraphs[0].add_run(text)
        run.font.name = '微软雅黑'
        run.font.size = Pt(10)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()

# ==================== 详细大纲 ====================
add_heading(doc, '二、详细大纲', level=1, size=16)

# 第1页
add_heading(doc, '第1页：封面', level=2, size=14)
add_bullet(doc, '主标题：AHL - AI驱动的去中心化旅行平台')
add_bullet(doc, '副标题：让每一程旅行都值得期待')
add_bullet(doc, '标签：AI + 旅行 + 去中心化')
doc.add_paragraph()

# 第2页
add_heading(doc, '第2页：我们在做什么', level=2, size=14)
add_bullet(doc, '标题：我们正在构建新一代旅行服务基础设施')
add_bullet(doc, '')
add_bullet(doc, '一句话定义：')
add_bullet(doc, 'AHL是一个AI驱动的去中心化旅行平台，让单体酒店和特色民宿也能拥有OTA级别的智能服务能力。', level=1)
add_bullet(doc, '')
add_bullet(doc, '做什么：')
add_bullet(doc, '为C端：AI管家提供7×24小时个性化旅行服务', level=1)
add_bullet(doc, '为B端：AI运营官帮助酒店提升收益30%+', level=1)
add_bullet(doc, '')
add_bullet(doc, '→ 简单说：我们用AI重新连接"旅行者"与"好酒店"')
doc.add_paragraph()

# 第3页
add_heading(doc, '第3页：我们要实现什么', level=2, size=14)
add_bullet(doc, '标题：6个月要做到的里程碑')
add_bullet(doc, '')
add_bullet(doc, '时间轴展示：')
add_bullet(doc, '第1个月：完成MVP核心功能开发', level=1)
add_bullet(doc, '第3个月：签约首批10家试点酒店', level=1)
add_bullet(doc, '第6个月：验证商业模式，覆盖50家酒店', level=1)
add_bullet(doc, '')
add_bullet(doc, '量化目标：')
add_bullet(doc, '帮助酒店提升收益30%+', level=1)
add_bullet(doc, '客户满意度提升20%+', level=1)
add_bullet(doc, 'AI回复准确率90%+', level=1)
add_bullet(doc, '')
add_bullet(doc, '→ 先用乐山锦江嘉州宾馆验证，再复制全国')
doc.add_paragraph()

# 第4页
add_heading(doc, '第4页：我们的愿景', level=2, size=14)
add_bullet(doc, '标题：让每一程旅行都值得期待')
add_bullet(doc, '')
add_bullet(doc, '我们相信：')
add_bullet(doc, '每个旅行者都值得被"懂你"的服务对待', level=1)
add_bullet(doc, '每家好酒店都值得被精准匹配到对的人', level=1)
add_bullet(doc, 'AI不是取代人，而是让人回归服务本质', level=1)
add_bullet(doc, '')
add_bullet(doc, '最终目标：打破OTA垄断，让酒店和旅行者直接对话')
doc.add_paragraph()

# 第5页
add_heading(doc, '第5页：行业痛点', level=2, size=14)
add_bullet(doc, '标题：传统OTA模式的四大顽疾')
add_bullet(doc, '')
add_bullet(doc, '1. 佣金高昂')
add_bullet(doc, '平台抽成15%-25%，酒店利润被侵蚀', level=1)
add_bullet(doc, '')
add_bullet(doc, '2. 数据不透明')
add_bullet(doc, '算法黑箱，酒店被动接受流量分配', level=1)
add_bullet(doc, '')
add_bullet(doc, '3. 用户资产缺失')
add_bullet(doc, '平台掌握用户数据，酒店无法转化', level=1)
add_bullet(doc, '')
add_bullet(doc, '4. 服务同质化')
add_bullet(doc, '陷入价格战，无法做差异化', level=1)
add_bullet(doc, '')
add_bullet(doc, '→ 行业亟需新范式')
doc.add_paragraph()

# 第6页
add_heading(doc, '第6页：解决方案', level=2, size=14)
add_bullet(doc, '标题：AHL双AGENT智能服务体系')
add_bullet(doc, '')
add_bullet(doc, 'C端AI管家：')
add_bullet(doc, '智能预订、个性化推荐、行程管家', level=1)
add_bullet(doc, '7×24小时服务，秒级响应', level=1)
add_bullet(doc, '记住偏好，提供"懂你"体验', level=1)
add_bullet(doc, '')
add_bullet(doc, 'B端AI运营官：')
add_bullet(doc, '动态定价、收益管理、OTA运营自动化', level=1)
add_bullet(doc, '盘活闲置资源，降低人工成本', level=1)
add_bullet(doc, '数据驱动决策，告别经验主义', level=1)
add_bullet(doc, '')
add_bullet(doc, '→ C端体验升级 + B端效率提升 = 双向赋能')
doc.add_paragraph()

# 第7页
add_heading(doc, '第7页：产品架构', level=2, size=14)
add_bullet(doc, '标题：AHL产品矩阵')
add_bullet(doc, '')
add_bullet(doc, '┌─────────────────────────────────┐')
add_bullet(doc, '│         AHL智能中枢              │')
add_bullet(doc, '├───────────────┬─────────────────┤')
add_bullet(doc, '│  C端AI管家    │   B端AI运营官   │')
add_bullet(doc, '├───────────────┼─────────────────┤')
add_bullet(doc, '│• 智能预订     │• 收益管理AGENT  │')
add_bullet(doc, '│• 个性化推荐   │• OTA运营AGENT   │')
add_bullet(doc, '│• 行程管家     │• 私域运营AGENT  │')
add_bullet(doc, '│• 投诉处理     │• 供应链AGENT    │')
add_bullet(doc, '└───────────────┴─────────────────┘')
add_bullet(doc, '')
add_bullet(doc, '特点：轻量化部署，3步接入现有PMS')
doc.add_paragraph()

# 第8页
add_heading(doc, '第8页：技术壁垒', level=2, size=14)
add_bullet(doc, '标题：我们的核心壁垒')
add_bullet(doc, '')
add_bullet(doc, '1. 私有知识库')
add_bullet(doc, '24年酒店行业Know-How结构化', level=1)
add_bullet(doc, '')
add_bullet(doc, '2. 场景化SKILL体系')
add_bullet(doc, '80+细分场景，像搭积木一样组合', level=1)
add_bullet(doc, '')
add_bullet(doc, '3. 双AGENT协同')
add_bullet(doc, 'C+B端数据互通，闭环优化', level=1)
add_bullet(doc, '')
add_bullet(doc, '4. 行业专属模型')
add_bullet(doc, '酒店行业语料训练，懂酒店语言', level=1)
add_bullet(doc, '')
add_bullet(doc, '→ 不是通用AI，是专为酒店训练的"行业专家"')
doc.add_paragraph()

# 第9页
add_heading(doc, '第9页：商业模式', level=2, size=14)
add_bullet(doc, '标题：多元收入，持续盈利')
add_bullet(doc, '')
add_bullet(doc, '1. SaaS订阅费')
add_bullet(doc, '按房间数/月收费，¥50-200/间/月', level=1)
add_bullet(doc, '')
add_bullet(doc, '2. 交易分成')
add_bullet(doc, 'GMV的2-5%服务费', level=1)
add_bullet(doc, '')
add_bullet(doc, '3. 增值服务')
add_bullet(doc, '私域运营、供应链推荐等', level=1)
add_bullet(doc, '')
add_bullet(doc, '4. 数据增值')
add_bullet(doc, '行业洞察报告（脱敏后）', level=1)
add_bullet(doc, '')
add_bullet(doc, '目标：第3年实现盈亏平衡')
doc.add_paragraph()

# 第10页
add_heading(doc, '第10页：市场机会', level=2, size=14)
add_bullet(doc, '标题：千亿级市场，AI渗透加速')
add_bullet(doc, '')
add_bullet(doc, '• 中国在线旅游市场：1.8万亿（2025E）')
add_bullet(doc, '• 单体酒店数量：60万家+')
add_bullet(doc, '• AI渗透率：<5%，蓝海窗口')
add_bullet(doc, '')
add_bullet(doc, '我们选择：单体精品酒店 + 特色民宿')
add_bullet(doc, '路径：西南地区 → 全国复制')
doc.add_paragraph()

# 第11页
add_heading(doc, '第11页：团队优势', level=2, size=14)
add_bullet(doc, '标题：为什么是我们')
add_bullet(doc, '')
add_bullet(doc, '张实 | 创始人')
add_bullet(doc, '• 24年酒店行业，深度理解全链条', level=1)
add_bullet(doc, '• 曾实现7年预测偏差<5%', level=1)
add_bullet(doc, '')
add_bullet(doc, 'AI团队')
add_bullet(doc, '• 核心成员来自头部大厂', level=1)
add_bullet(doc, '• 酒店行业Know-How持续输入', level=1)
add_bullet(doc, '')
add_bullet(doc, '→ 行业洞察 × AI技术 = 真壁垒')
doc.add_paragraph()

# 第12页
add_heading(doc, '第12页：融资需求', level=2, size=14)
add_bullet(doc, '标题：种子轮融资计划')
add_bullet(doc, '')
add_bullet(doc, '融资规模：500-800万')
add_bullet(doc, '')
add_bullet(doc, '资金用途：')
add_bullet(doc, '40% 技术研发', level=1)
add_bullet(doc, '30% 市场拓展', level=1)
add_bullet(doc, '20% 团队扩充', level=1)
add_bullet(doc, '10% 运营储备', level=1)
add_bullet(doc, '')
add_bullet(doc, '里程碑：6个月MVP / 12个月50家酒店')
doc.add_paragraph()

# ==================== 附录 ====================
add_heading(doc, '三、OPC路演注意事项', level=1, size=16)

table2 = doc.add_table(rows=1, cols=2)
table2.style = 'Table Grid'

hdr_cells = table2.rows[0].cells
for i, text in enumerate(['要点', '说明']):
    hdr_cells[i].paragraphs[0].clear()
    run = hdr_cells[i].paragraphs[0].add_run(text)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run.font.size = Pt(10)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

tips = [
    ('语速', '控制在每页30-40秒，平均语速'),
    ('开场', '直接点题，不要寒暄'),
    ('痛点', '用数据说话，要有冲击力'),
    ('产品', '画图比说话更直观'),
    ('团队', '突出差异化优势'),
    ('融资', '简洁明了，数据支撑'),
    ('Q&A', '预判3个问题：技术壁垒、商业模式、竞争格局'),
]

for tip_data in tips:
    row = table2.add_row()
    for i, text in enumerate(tip_data):
        cell = row.cells[i]
        run = cell.paragraphs[0].add_run(text)
        run.font.name = '微软雅黑'
        run.font.size = Pt(10)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()

# ==================== 底部 ====================
add_heading(doc, '四、调整要点', level=1, size=16)

add_bullet(doc, '原结构：先痛点 → 再愿景')
add_bullet(doc, '调整后：先说做什么 → 再说要实现什么 → 然后愿景 → 最后痛点')
add_bullet(doc, '')
add_bullet(doc, '逻辑线：我们在做什么 → 要实现什么 → 我们相信什么 → 行业问题 → 我们的解法')
add_bullet(doc, '')
add_bullet(doc, '核心变化：')
add_bullet(doc, '开门见山，让听众第一时间清楚AHL是做什么的', level=1)
add_bullet(doc, '愿景单独成页，强化信念感', level=1)
add_bullet(doc, '痛点后置，建立在"我们能解决"的自信基础上', level=1)

# 保存
output_path = r'C:\Users\ericz\Desktop\商业计划书\AHL-OPC路演PPT大纲-精华版.docx'
doc.save(output_path)
print(f"Document saved to: {output_path}")
